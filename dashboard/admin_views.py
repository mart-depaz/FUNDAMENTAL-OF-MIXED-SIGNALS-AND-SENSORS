# dashboard/admin_views.py
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.http import JsonResponse, HttpResponse
from django.views.decorators.http import require_http_methods
from django.core.mail import send_mail
from django.template.loader import render_to_string
from django.conf import settings
from django.db import models
from django.utils import timezone
from accounts.models import CustomUser
from accounts.admin_forms import AdminAddTeacherForm, AdminAddStudentForm
from .models import Program, AdminNotification, Course, Department, CourseSchedule
import json
import logging
import random
from io import BytesIO
from datetime import datetime
import csv
from zoneinfo import ZoneInfo

PH_TIMEZONE = ZoneInfo("Asia/Manila")

logger = logging.getLogger(__name__)


def format_year_label(year_level):
    try:
        year = int(year_level)
    except (ValueError, TypeError):
        return f"Year {year_level}"
    suffix = "th"
    if 10 <= year % 100 <= 20:
        suffix = "th"
    else:
        if year % 10 == 1:
            suffix = "st"
        elif year % 10 == 2:
            suffix = "nd"
        elif year % 10 == 3:
            suffix = "rd"
    return f"{year}{suffix} Year"

# Try to import python-docx, but don't fail if it's not available
# This is an optional dependency - the system will work without it
DOCX_AVAILABLE = False
try:
    import docx
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_SECTION
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    DOCX_AVAILABLE = True
except (ImportError, ModuleNotFoundError):
    # Silently handle missing python-docx - it's optional
    # The system will work without it, just Word document generation won't be available
    pass

def generate_user_document(user, temp_password, user_type='instructor', login_url=None):
    """Generate a Word document with user information - optimized for printing"""
    if not DOCX_AVAILABLE:
        return None
    
    # Import here to avoid issues during module load if docx is not available
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return None
    
    # Ensure password is always shown (not a placeholder)
    if not temp_password or temp_password.startswith('********'):
        # This shouldn't happen if called correctly, but handle it gracefully
        temp_password = 'Please contact administrator for password'
    
    doc = Document()
    
    # Set up page for printing (Letter size, standard margins)
    section = doc.sections[0]
    section.page_height = Inches(11)  # Letter height
    section.page_width = Inches(8.5)    # Letter width
    section.left_margin = Inches(1)     # 1 inch left margin
    section.right_margin = Inches(1)    # 1 inch right margin
    section.top_margin = Inches(1)      # 1 inch top margin
    section.bottom_margin = Inches(1)   # 1 inch bottom margin
    
    # Title
    title_text = f'Account Information - {"Instructor" if user_type == "instructor" else "Student"}'
    title = doc.add_heading(title_text, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # School Name Header
    if user.school_name:
        school_para = doc.add_paragraph()
        school_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        school_run = school_para.add_run(user.school_name.upper())
        school_run.font.size = Pt(16)
        school_run.font.bold = True
        school_run.font.color.rgb = RGBColor(0, 0, 0)  # Black for better printing
    
    doc.add_paragraph()  # Spacing
    
    # Date
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_run = date_para.add_run(f'Date: {datetime.now().strftime("%B %d, %Y")}')
    date_run.font.size = Pt(10)
    date_run.italic = True
    
    doc.add_paragraph()  # Spacing
    
    # Personal Information Section
    doc.add_heading('Personal Information', level=1)
    
    info_data = [
        ('Full Name', user.full_name or 'N/A'),
        ('Email Address', user.email or 'N/A'),
        ('ID Number', user.school_id or 'N/A'),
        ('Username', user.username or 'N/A'),
        ('Education Level', user.get_education_level_display() if user.education_level else 'N/A'),
        ('School', user.school_name or 'N/A'),
    ]
    
    if user.department:
        info_data.append(('Department', user.department))
    
    if user.program:
        info_data.append(('Program', f"{user.program.code} - {user.program.name}"))
    
    if user_type == 'student':
        if user.year_level:
            year_text = f"{user.year_level}{'st' if user.year_level == 1 else 'nd' if user.year_level == 2 else 'rd' if user.year_level == 3 else 'th'} Year"
            info_data.append(('Year Level', year_text))
        if user.section:
            info_data.append(('Section', user.section))
    
    # Create table for information
    table = doc.add_table(rows=len(info_data), cols=2)
    # Use a basic table style (will work even if specific style doesn't exist)
    try:
        table.style = 'Light Grid Accent 1'
    except:
        pass  # Use default style if custom style not available
    
    # Set table column widths for better printing
    for row in table.rows:
        row.cells[0].width = Inches(2.5)  # Label column
        row.cells[1].width = Inches(4.5)   # Value column
    
    for i, (label, value) in enumerate(info_data):
        # Clear and set label with styling
        table.rows[i].cells[0].paragraphs[0].clear()
        label_run = table.rows[i].cells[0].paragraphs[0].add_run(label)
        label_run.font.bold = True
        label_run.font.color.rgb = RGBColor(0, 0, 0)  # Black for better printing
        
        # Set value with readable font
        table.rows[i].cells[1].text = str(value)
        if table.rows[i].cells[1].paragraphs[0].runs:
            table.rows[i].cells[1].paragraphs[0].runs[0].font.size = Pt(11)
    
    doc.add_paragraph()  # Spacing
    
    # Login Credentials Section
    doc.add_heading('Login Credentials', level=1)
    
    cred_table = doc.add_table(rows=3, cols=2)
    try:
        cred_table.style = 'Light Grid Accent 1'
    except:
        pass  # Use default style if custom style not available
    
    # Set table column widths for better printing
    for row in cred_table.rows:
        row.cells[0].width = Inches(2.5)  # Label column
        row.cells[1].width = Inches(4.5)   # Value column
    
    # Build login URL
    if not login_url:
        login_url = 'Please contact your administrator for the login URL'
    
    cred_data = [
        ('Login Email/ID', f"{user.email} or {user.school_id}"),
        ('Temporary Password', temp_password),  # Always show the actual password
        ('Login URL', login_url)
    ]
    
    for i, (label, value) in enumerate(cred_data):
        # Clear and set label with styling
        cred_table.rows[i].cells[0].paragraphs[0].clear()
        label_run = cred_table.rows[i].cells[0].paragraphs[0].add_run(label)
        label_run.font.bold = True
        label_run.font.color.rgb = RGBColor(0, 0, 0)  # Black for better printing
        
        # Set value with special styling for password
        cred_table.rows[i].cells[1].paragraphs[0].clear()
        value_run = cred_table.rows[i].cells[1].paragraphs[0].add_run(str(value))
        value_run.font.size = Pt(11)
        if label == 'Temporary Password':
            # Make password very prominent - bold, larger font
            value_run.font.bold = True
            value_run.font.size = Pt(14)  # Larger font for password visibility
            value_run.font.color.rgb = RGBColor(0, 0, 0)  # Black for printing
            # Add helpful note below the password
            note_run = cred_table.rows[i].cells[1].paragraphs[0].add_run('\n(This is your login password - keep it safe)')
            note_run.font.size = Pt(9)
            note_run.font.italic = True
            note_run.font.color.rgb = RGBColor(128, 128, 128)
    
    doc.add_paragraph()  # Spacing
    
    # Important Notes
    doc.add_heading('Important Notes', level=1)
    notes = [
        "Please keep this document in a safe place.",
        "Change your password immediately after first login for security.",
        "Do not share your login credentials with anyone.",
        "If you forget your password, contact your school administrator.",
    ]
    
    for note in notes:
        para = doc.add_paragraph(note, style='List Bullet')
        para.runs[0].font.size = Pt(11)
    
    doc.add_paragraph()  # Spacing
    
    # Footer
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run('Generated by Attendance System')
    footer_run.font.size = Pt(9)
    footer_run.italic = True
    footer_run.font.color.rgb = RGBColor(128, 128, 128)
    
    # Save to BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer

def admin_required(view_func):
    """
    Decorator to ensure only school admins can access a view.
    Note: Django superusers (is_superuser=True) should use Django admin at /admin/
    School admins (is_admin=True) use this custom portal at /admin-portal/
    """
    def wrapper(request, *args, **kwargs):
        # Check if this is an AJAX/API request (JSON, FormData, or XMLHttpRequest)
        content_type = request.META.get('CONTENT_TYPE', '') or ''
        is_ajax = (
            request.headers.get('X-Requested-With') == 'XMLHttpRequest' or
            request.content_type == 'application/json' or
            'multipart/form-data' in content_type or
            request.method == 'POST' and (content_type.startswith('multipart/') or content_type.startswith('application/json'))
        )
        
        if not request.user.is_authenticated:
            # For AJAX requests, return JSON instead of redirect
            if is_ajax:
                from django.http import JsonResponse
                return JsonResponse({'success': False, 'message': 'Authentication required. Please log in.'}, status=401)
            return redirect('admin_login_signup')
        # Only allow school admins (not Django superusers accessing this portal)
        if not request.user.is_admin:
            # For AJAX requests, return JSON instead of render
            if is_ajax:
                from django.http import JsonResponse
                return JsonResponse({'success': False, 'message': 'You are not authorized to access this resource.'}, status=403)
            return render(request, 'dashboard/shared/error.html', {
                'message': 'You are not authorized to access the school admin dashboard. This portal is for school administrators only. Django superusers should use /admin/ instead.'
            })
        return view_func(request, *args, **kwargs)
    return wrapper

@login_required
@admin_required
def admin_dashboard_view(request):
    """Main admin dashboard view"""
    user = request.user
    
    # Note: Messages are handled by the notification system
    # Logout messages are filtered out on dashboard pages by the notification system
    
    # Get statistics - filtered by school, excluding deleted items
    if user.school_name:
        total_departments = Department.objects.filter(school_name=user.school_name, is_active=True, deleted_at__isnull=True).count()
        total_programs = Program.objects.filter(school_name=user.school_name, deleted_at__isnull=True).count()
        active_programs = Program.objects.filter(school_name=user.school_name, is_active=True, deleted_at__isnull=True).count()
        # Grouped total courses (one per code+name+semester+school_year)
        courses_qs = Course.objects.filter(school_name=user.school_name)
        course_keys = set()
        for c in courses_qs.values('code', 'name', 'semester', 'school_year'):
            norm_code = (c.get('code') or '').strip().upper()
            norm_name = (c.get('name') or '').strip().upper()
            norm_sem = (c.get('semester') or '').strip().lower()
            norm_sy = (c.get('school_year') or '').strip()
            course_keys.add((norm_code, norm_name, norm_sem, norm_sy))
        total_courses = len(course_keys)
        total_teachers = CustomUser.objects.filter(is_teacher=True, school_name=user.school_name, is_approved=True, deleted_at__isnull=True).count()
        total_students = CustomUser.objects.filter(is_student=True, school_name=user.school_name, is_approved=True, deleted_at__isnull=True).count()  # Only approved students
        active_users = CustomUser.objects.filter(
            school_name=user.school_name,
            is_approved=True,
            deleted_at__isnull=True
        ).exclude(is_admin=True).count()  # Active users (teachers + students, excluding admins)
        # Only teachers need approval (students are auto-approved)
        pending_approvals = CustomUser.objects.filter(
            is_teacher=True,
            is_approved=False,
            school_name=user.school_name,
            deleted_at__isnull=True
        ).count()
    else:
        total_departments = Department.objects.filter(is_active=True, deleted_at__isnull=True).count()
        total_programs = Program.objects.filter(deleted_at__isnull=True).count()
        active_programs = Program.objects.filter(is_active=True, deleted_at__isnull=True).count()
        # Grouped total courses (one per code+name+semester+school_year)
        courses_qs = Course.objects.all()
        course_keys = set()
        for c in courses_qs.values('code', 'name', 'semester', 'school_year'):
            norm_code = (c.get('code') or '').strip().upper()
            norm_name = (c.get('name') or '').strip().upper()
            norm_sem = (c.get('semester') or '').strip().lower()
            norm_sy = (c.get('school_year') or '').strip()
            course_keys.add((norm_code, norm_name, norm_sem, norm_sy))
        total_courses = len(course_keys)
        total_teachers = CustomUser.objects.filter(is_teacher=True, is_approved=True, deleted_at__isnull=True).count()
        total_students = CustomUser.objects.filter(is_student=True, is_approved=True, deleted_at__isnull=True).count()  # Only approved students
        active_users = CustomUser.objects.filter(
            is_approved=True,
            deleted_at__isnull=True
        ).exclude(is_admin=True).count()  # Active users (teachers + students, excluding admins)
        # Only teachers need approval (students are auto-approved)
        pending_approvals = CustomUser.objects.filter(is_teacher=True, is_approved=False, deleted_at__isnull=True).count()
    
    # Get unread notifications
    unread_notifications = AdminNotification.objects.filter(admin=user, is_read=False).count()
    
    context = {
        'user': user,
        'total_departments': total_departments,
        'total_programs': total_programs,
        'active_programs': active_programs,
        'total_courses': total_courses,
        'total_teachers': total_teachers,
        'total_students': total_students,
        'active_users': active_users,
        'pending_approvals': pending_approvals,
        'unread_notifications': unread_notifications,
    }
    
    return render(request, 'dashboard/admin/admin_dashboard.html', context)

@login_required
@admin_required
def admin_institutional_setup_view(request):
    """Institutional setup view for managing programs with search"""
    user = request.user
    
    # Get search query
    search_query = request.GET.get('search', '').strip()
    
    # Get programs filtered by school and education level (exclude deleted)
    if user.school_name:
        programs = Program.objects.filter(school_name=user.school_name, deleted_at__isnull=True).order_by('code')
    else:
        programs = Program.objects.filter(deleted_at__isnull=True).order_by('code')
    
    # Filter by admin's education level - only show programs matching their level
    if user.education_level:
        programs = programs.filter(education_level=user.education_level)
    
    # Get department filter
    department_filter = request.GET.get('department_filter', '').strip()
    
    # Apply department filter if provided
    if department_filter:
        programs = programs.filter(department=department_filter)
    
    # Apply search filter if provided
    if search_query:
        programs = programs.filter(
            models.Q(code__icontains=search_query) |
            models.Q(name__icontains=search_query) |
            models.Q(department__icontains=search_query)
        )
    
    # Get user counts for each program (exclude deleted users)
    program_data = []
    for program in programs:
        if user.school_name:
            users = CustomUser.objects.filter(program=program, school_name=user.school_name, deleted_at__isnull=True)
        else:
            users = CustomUser.objects.filter(program=program, deleted_at__isnull=True)
        program_data.append({
            'program': program,
            'user_count': users.count(),
            'teachers_count': users.filter(is_teacher=True).count(),
            'students_count': users.filter(is_student=True).count(),
        })
    
    # Group programs by department
    programs_by_department = {}
    for pd in program_data:
        dept = pd['program'].department or 'Uncategorized'
        if dept not in programs_by_department:
            programs_by_department[dept] = []
        programs_by_department[dept].append(pd)
    
    unread_notifications = AdminNotification.objects.filter(admin=user, is_read=False).count()
    
    # Get all programs for program filter (grouped by department) - exclude deleted
    all_programs_by_dept = {}
    if user.school_name:
        all_programs = Program.objects.filter(school_name=user.school_name, deleted_at__isnull=True).order_by('code')
    else:
        all_programs = Program.objects.filter(deleted_at__isnull=True).order_by('code')
    
    if user.education_level:
        all_programs = all_programs.filter(education_level=user.education_level)
    
    for program in all_programs:
        dept = program.department or 'Uncategorized'
        if dept not in all_programs_by_dept:
            all_programs_by_dept[dept] = []
        all_programs_by_dept[dept].append(program)
    
    # Get all departments - newest appears right after the previous added
    if user.school_name:
        all_departments = Department.objects.filter(school_name=user.school_name, is_active=True, deleted_at__isnull=True).order_by('created_at', 'id')
    else:
        all_departments = Department.objects.filter(is_active=True, deleted_at__isnull=True).order_by('created_at', 'id')
    
    if user.education_level:
        all_departments = all_departments.filter(education_level=user.education_level)
    
    # Convert to list - newest is already at the end, which is correct
    all_departments = list(all_departments)
    
    # Group programs by Department object (not just department name)
    departments_with_programs = []
    for dept in all_departments:
        dept_programs = []
        for pd in program_data:
            # Match by department name or department code
            if pd['program'].department == dept.name or pd['program'].department_code == dept.code:
                dept_programs.append(pd)
        
        departments_with_programs.append({
            'department': dept,
            'programs': dept_programs,
            'program_count': len(dept_programs)
        })
    
    # Get department filter for search
    department_filter = request.GET.get('department_filter', '').strip()
    selected_department_id = request.GET.get('dept', '').strip()
    
    context = {
        'user': user,
        'programs': programs,
        'program_data': program_data,
        'programs_by_department': programs_by_department,
        'all_programs_by_dept': all_programs_by_dept,
        'departments': all_departments,
        'departments_with_programs': departments_with_programs,  # New structure
        'selected_department_id': selected_department_id,
        'education_level': user.education_level,
        'unread_notifications': unread_notifications,
        'search_query': search_query,
    }
    
    return render(request, 'dashboard/admin/admin_institutional_setup.html', context)

@login_required
@admin_required
def admin_department_programs_view(request, department_id):
    """View all programs in a specific department"""
    user = request.user
    department = get_object_or_404(Department, id=department_id)
    
    # Verify department belongs to admin's school
    if user.school_name and department.school_name != user.school_name:
        return render(request, 'dashboard/shared/error.html', {
            'message': 'You are not authorized to view this department.'
        })
    
    # Get all programs in this department (exclude deleted) - newest appears right after the previous added
    if user.school_name:
        programs = Program.objects.filter(
            department=department.name,
            school_name=user.school_name,
            is_active=True,
            deleted_at__isnull=True
        ).order_by('created_at', 'id')
    else:
        programs = Program.objects.filter(
            department=department.name,
            is_active=True,
            deleted_at__isnull=True
        ).order_by('created_at', 'id')
    
    # Convert to list - newest is already at the end, which is correct
    programs = list(programs)
    
    # Get program data with user counts (exclude deleted users)
    program_data = []
    for program in programs:
        if user.school_name:
            user_count = CustomUser.objects.filter(program=program, school_name=user.school_name, deleted_at__isnull=True).count()
        else:
            user_count = CustomUser.objects.filter(program=program, deleted_at__isnull=True).count()
        
        program_data.append({
            'program': program,
            'user_count': user_count
        })
    
    unread_notifications = AdminNotification.objects.filter(admin=user, is_read=False).count()
    
    # Get all departments for navigation (exclude deleted) - newest appears right after the previous added
    if user.school_name:
        all_departments = Department.objects.filter(school_name=user.school_name, is_active=True, deleted_at__isnull=True).order_by('created_at', 'id')
    else:
        all_departments = Department.objects.filter(is_active=True, deleted_at__isnull=True).order_by('created_at', 'id')
    
    # Apply education level filter BEFORE converting to list
    if user.education_level:
        all_departments = all_departments.filter(education_level=user.education_level)
    
    # Convert to list - newest is already at the end, which is correct
    all_departments = list(all_departments)
    
    context = {
        'user': user,
        'department': department,
        'program_data': program_data,
        'all_departments': all_departments,
        'unread_notifications': unread_notifications,
    }
    
    return render(request, 'dashboard/admin/admin_department_programs.html', context)

@login_required
@admin_required
def admin_program_users_view(request, program_id):
    """View users enrolled in a specific program with search and filter"""
    user = request.user
    program = get_object_or_404(Program, id=program_id)
    
    # Verify program belongs to admin's school
    if user.school_name and program.school_name != user.school_name:
        return render(request, 'dashboard/shared/error.html', {
            'message': 'You are not authorized to view this program.'
        })
    
    # Get search query and filters
    search_query = request.GET.get('search', '').strip()
    section_filter = request.GET.get('section_filter', '').strip()
    year_level_filter = request.GET.get('year_level_filter', '').strip()
    school_year_filter = request.GET.get('school_year_filter', '').strip()
    user_has_school_year = hasattr(CustomUser, 'school_year')
    
    # Get users in this program (exclude deleted)
    if user.school_name:
        users = CustomUser.objects.filter(program=program, school_name=user.school_name, deleted_at__isnull=True)
    else:
        users = CustomUser.objects.filter(program=program, deleted_at__isnull=True)
    
    # Apply section filter (for students)
    if section_filter:
        users = users.filter(section=section_filter)
    
    # Apply year level filter (for students)
    if year_level_filter:
        try:
            year_level = int(year_level_filter)
            users = users.filter(year_level=year_level)
        except ValueError:
            pass
    
    if school_year_filter and user_has_school_year:
        try:
            users = users.filter(school_year=school_year_filter)
        except Exception:
            school_year_filter = ''
    
    # Apply search filter
    if search_query:
        users = users.filter(
            models.Q(full_name__icontains=search_query) |
            models.Q(email__icontains=search_query) |
            models.Q(school_id__icontains=search_query) |
            models.Q(username__icontains=search_query) |
            models.Q(department__icontains=search_query)
        )
    
    # Order by creation date - newest appears right after the previous added
    users = users.order_by('date_joined', 'id')
    teachers = users.filter(is_teacher=True)
    students = users.filter(is_student=True)
    
    # Get unique sections and year levels for filter dropdown (from students in this program)
    if user.school_name:
        all_students = CustomUser.objects.filter(program=program, school_name=user.school_name, is_student=True)
    else:
        all_students = CustomUser.objects.filter(program=program, is_student=True)
    unique_sections = sorted(set([s.section for s in all_students if s.section]))
    unique_year_levels = sorted(set([s.year_level for s in all_students if s.year_level]))
    if user_has_school_year:
        try:
            unique_school_years = sorted(
                filter(
                    None,
                    users.values_list('school_year', flat=True).distinct()
                ),
                reverse=True
            )
        except Exception:
            unique_school_years = []
            school_year_filter = ''
    else:
        unique_school_years = []
        school_year_filter = ''
    
    unread_notifications = AdminNotification.objects.filter(admin=user, is_read=False).count()
    
    # Get all programs for the add user modals (filtered by school and education level)
    if user.school_name:
        all_programs = Program.objects.filter(school_name=user.school_name).order_by('code')
    else:
        all_programs = Program.objects.all().order_by('code')
    
    if user.education_level:
        all_programs = all_programs.filter(education_level=user.education_level)
    
    context = {
        'user': user,
        'program': program,
        'users': users,
        'teachers': teachers,
        'students': students,
        'programs': all_programs,  # For modal dropdowns
        'unread_notifications': unread_notifications,
        'search_query': search_query,
        'section_filter': section_filter,
        'year_level_filter': year_level_filter,
        'school_year_filter': school_year_filter,
        'unique_sections': unique_sections,
        'unique_year_levels': unique_year_levels,
        'unique_school_years': unique_school_years,
        'from_database': request.GET.get('from') == 'database',
    }
    
    return render(request, 'dashboard/admin/admin_program_users.html', context)


@login_required
@admin_required
def admin_download_users_csv_view(request):
    """Allow admins to download user data (all, by program, or individual)."""
    user = request.user
    program_id = request.GET.get('program_id')
    user_id = request.GET.get('user_id')
    search_query = request.GET.get('search', '').strip()
    section_filter = request.GET.get('section_filter', request.GET.get('section', '')).strip()
    year_level_filter = request.GET.get('year_level_filter', request.GET.get('year_level', '')).strip()
    user_type = request.GET.get('user_type', 'all').strip().lower()

    queryset = CustomUser.objects.filter(deleted_at__isnull=True)

    if user.school_name:
        queryset = queryset.filter(school_name=user.school_name)

    if program_id:
        try:
            program = Program.objects.get(id=program_id)
            if user.school_name and program.school_name != user.school_name:
                return JsonResponse({'success': False, 'message': 'You are not authorized to download these users.'}, status=403)
        except Program.DoesNotExist:
            return JsonResponse({'success': False, 'message': 'Program not found.'}, status=404)
        queryset = queryset.filter(program_id=program_id)

    if user_id:
        queryset = queryset.filter(id=user_id)

    if section_filter:
        queryset = queryset.filter(section__iexact=section_filter)

    if year_level_filter:
        try:
            queryset = queryset.filter(year_level=int(year_level_filter))
        except ValueError:
            pass

    if search_query:
        queryset = queryset.filter(
            models.Q(full_name__icontains=search_query) |
            models.Q(email__icontains=search_query) |
            models.Q(school_id__icontains=search_query) |
            models.Q(username__icontains=search_query) |
            models.Q(department__icontains=search_query)
        )

    if user_type == 'instructors':
        queryset = queryset.filter(is_teacher=True)
    elif user_type == 'students':
        queryset = queryset.filter(is_student=True)
    else:
        user_type = 'all'
    
    if not queryset.exists():
        return JsonResponse({'success': False, 'message': 'No users found for export.'}, status=404)

    timestamp = datetime.now().strftime('%Y%m%d-%H%M%S')
    if user_id:
        filename = f'user-{user_id}-{timestamp}.csv'
    elif program_id:
        filename = f'program-{program_id}-users-{timestamp}.csv'
    else:
        filename = f'all-users-{timestamp}.csv'

    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = f'attachment; filename="{filename}"'

    file_format = request.GET.get('format', 'csv').lower()
    
    def get_headers(include_student_fields=True):
        base_headers = [
            'User ID', 'Full Name', 'Email', 'School ID',
            'User Type', 'Program', 'Department'
        ]
        if include_student_fields:
            base_headers.extend(['Year Level', 'Section'])
        base_headers.extend(['Status', 'Date Joined'])
        return base_headers
    
    def format_row(entry, include_student_fields=True, for_csv=False):
        role = 'Instructor' if entry.is_teacher else 'Student' if entry.is_student else 'User'
        status = 'Active' if entry.is_approved else 'Pending'
        dt = entry.date_joined
        if dt:
            if timezone.is_naive(dt):
                dt = timezone.make_aware(dt, timezone.get_default_timezone())
            dt = timezone.localtime(dt, PH_TIMEZONE)
            date_str = dt.strftime('%Y-%m-%d %I:%M %p')
            if for_csv:
                date_str = f"'{date_str}"
        else:
            date_str = ''
        row = [
            entry.id,
            entry.full_name or entry.username,
            entry.email,
            entry.school_id or '',
            role,
            entry.program.code if entry.program else '',
            entry.department or '',
        ]
        if include_student_fields:
            row.extend([
                entry.year_level or '',
                entry.section or '',
            ])
        row.extend([status, date_str])
        return row
    
    if file_format == 'docx':
        if not DOCX_AVAILABLE:
            return HttpResponse('Word export not available. Please install python-docx.', status=500)
        doc = Document()
        header_table = doc.add_table(rows=1, cols=2)
        header_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        header_table.autofit = False
        header_table.columns[0].width = Inches(1.0)
        header_table.columns[1].width = Inches(5.5)
        school_image = getattr(user, 'profile_picture', None)
        left_cell = header_table.cell(0, 0)
        left_cell.text = ''
        left_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        if school_image and hasattr(school_image, 'path'):
            try:
                left_cell.paragraphs[0].add_run().add_picture(school_image.path, width=Inches(0.9))
            except Exception:
                pass
        right_cell = header_table.cell(0, 1)
        right_cell.text = ''
        right_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        name_para = right_cell.paragraphs[0]
        name_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        name_run = name_para.add_run(user.school_name or 'School')
        name_run.font.bold = True
        name_run.font.size = Pt(12)
        doc.add_paragraph()
        if program_id and 'program' in locals():
            program_table = doc.add_table(rows=1, cols=2)
            program_table.alignment = WD_TABLE_ALIGNMENT.LEFT
            program_table.autofit = False
            program_table.columns[0].width = Inches(1.0)
            program_table.columns[1].width = Inches(5.5)
            program_image = getattr(program, 'icon', None)
            p_left = program_table.cell(0, 0)
            p_left.text = ''
            p_left.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if program_image and hasattr(program_image, 'path'):
                try:
                    p_left.paragraphs[0].add_run().add_picture(program_image.path, width=Inches(0.9))
                except Exception:
                    pass
            p_right = program_table.cell(0, 1)
            p_right.text = ''
            p_right.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p_para = p_right.paragraphs[0]
            p_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p_run = p_para.add_run(f"{program.code} - {program.name} ({program.department})")
            p_run.font.bold = True
            p_run.font.size = Pt(10)
            underline_para = doc.add_paragraph()
            underline_run = underline_para.add_run('_' * 90)
            underline_run.font.size = Pt(6)
        doc.add_paragraph(f'Generated: {datetime.now(PH_TIMEZONE).strftime("%B %d, %Y %I:%M %p")} (Philippine Time)')
        doc.add_paragraph()
        
        def add_doc_table(title, entries):
            if not entries:
                return False
            doc.add_heading(title, level=1)
            include_student_fields = (title != 'Instructors')
            headers = get_headers(include_student_fields)
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.LEFT
            table.autofit = False

            width_presets = {
                'User ID': 1.2,
                'Full Name': 3.8,
                'Email': 4.2,
                'School ID': 1.5,
                'User Type': 1.1,
                'Program': 1.8,
                'Department': 3.8,
                'Year Level': 1.3,
                'Section': 1.3,
                'Status': 1.2,
                'Date Joined': 2.0
            }
            default_width = 1.4

            hdr_cells = table.rows[0].cells
            for idx, header in enumerate(headers):
                cell = hdr_cells[idx]
                cell.text = ''
                paragraph = cell.paragraphs[0]
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run(header)
                run.bold = True
                run.font.size = Pt(7.5)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                cell.width = Inches(width_presets.get(header, default_width))

            for entry in entries:
                row_cells = table.add_row().cells
                row_values = format_row(entry, include_student_fields)
                for idx, value in enumerate(row_values):
                    cell = row_cells[idx]
                    cell.text = ''
                    paragraph = cell.paragraphs[0]
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    run = paragraph.add_run(str(value))
                    run.font.size = Pt(7.5)
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    cell.width = Inches(width_presets.get(headers[idx], default_width))

            doc.add_paragraph()
            return True
        
        if user_type == 'all':
            instructors = list(queryset.filter(is_teacher=True).order_by('full_name', 'username'))
            students = list(queryset.filter(is_student=True).order_by('full_name', 'username'))
            if instructors:
                add_doc_table('Instructors', instructors)
            if students:
                add_doc_table('Students', students)
        elif user_type == 'instructors':
            add_doc_table('Instructors', list(queryset.order_by('full_name', 'username')))
        else:
            add_doc_table('Students', list(queryset.order_by('full_name', 'username')))
        
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        filename_prefix = 'users'
        if user_type == 'instructors':
            filename_prefix = 'instructors'
        elif user_type == 'students':
            filename_prefix = 'students'
        response = HttpResponse(
            buffer.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="{filename_prefix}-{timestamp}.docx"'
        return response
    
    writer = csv.writer(response)
    if user_type == 'all':
        groups = [
            ('Instructors', list(queryset.filter(is_teacher=True).order_by('full_name', 'username'))),
            ('Students', list(queryset.filter(is_student=True).order_by('full_name', 'username')))
        ]
        wrote_section = False
        for title, entries in groups:
            if not entries:
                continue
            if wrote_section:
                writer.writerow([])
            writer.writerow([title])
            include_student_fields_csv = (title != 'Instructors')
            headers = get_headers(include_student_fields_csv)
            writer.writerow(headers)
            for entry in entries:
                writer.writerow(format_row(entry, include_student_fields=include_student_fields_csv, for_csv=True))
            wrote_section = True
        if not wrote_section:
            headers = get_headers(True)
            writer.writerow(headers)
    else:
        include_student_fields_csv = (user_type != 'instructors')
        headers = get_headers(include_student_fields_csv)
        writer.writerow(headers)
        for entry in queryset.order_by('full_name', 'username'):
            writer.writerow(format_row(entry, include_student_fields=include_student_fields_csv, for_csv=True))

    return response

@login_required
@admin_required
@require_http_methods(["POST"])
def admin_add_program_view(request):
    """Add a new program"""
    user = request.user
    
    try:
        # Handle both JSON and form data
        if request.content_type == 'application/json':
            data = json.loads(request.body)
            code = data.get('code', '').strip().upper()
            name = data.get('name', '').strip()
            department_id = data.get('department_id', '').strip()
        else:
            code = request.POST.get('code', '').strip().upper()
            name = request.POST.get('name', '').strip()
            department_id = request.POST.get('department_id', '').strip()
        
        if not code or not name or not department_id:
            return JsonResponse({'success': False, 'message': 'All fields are required.'})
        
        # Get the department
        try:
            department = Department.objects.get(id=department_id, school_name=user.school_name)
        except Department.DoesNotExist:
            return JsonResponse({'success': False, 'message': 'Invalid department selected.'})
        
        # Check if program with same code exists in this school
        existing = Program.objects.filter(code=code, school_name=user.school_name)
        if existing.exists():
            return JsonResponse({'success': False, 'message': f'Program with code {code} already exists in your school.'})
        
        # Handle icon upload
        icon = None
        if 'icon' in request.FILES:
            icon = request.FILES['icon']
        
        program = Program.objects.create(
            code=code,
            name=name,
            department=department.name,
            department_code=department.code,
            icon=icon,
            school_name=user.school_name,
            education_level=user.education_level or 'university_college',
            is_active=True
        )
        
        return JsonResponse({
            'success': True,
            'message': 'Program created successfully.',
            'program': {
                'id': program.id,
                'code': program.code,
                'name': program.name,
                'department': program.department,
            }
        })
    except Exception as e:
        return JsonResponse({'success': False, 'message': str(e)})

@login_required
@admin_required
@require_http_methods(["POST"])
def admin_update_program_view(request, program_id):
    """Update an existing program"""
    user = request.user
    
    try:
        program = get_object_or_404(Program, id=program_id)
        
        # Verify the program belongs to the admin's school
        if program.school_name and program.school_name != user.school_name:
            return JsonResponse({'success': False, 'message': 'You do not have permission to edit this program.'})
        
        # Handle both JSON and form data
        if request.content_type == 'application/json':
            data = json.loads(request.body)
            program.code = data.get('code', program.code).strip().upper()
            program.name = data.get('name', program.name).strip()
            department_id = data.get('department_id', '')
        else:
            program.code = request.POST.get('code', program.code).strip().upper()
            program.name = request.POST.get('name', program.name).strip()
            department_id = request.POST.get('department_id', '')
            
            # Handle icon upload
            if 'icon' in request.FILES:
                program.icon = request.FILES['icon']
        
        # Update department if department_id is provided
        if department_id:
            try:
                department = Department.objects.get(id=department_id, school_name=user.school_name)
                program.department = department.name
                program.department_code = department.code
            except Department.DoesNotExist:
                pass  # Keep existing department if invalid
        
        # Check for duplicate code in same school
        existing = Program.objects.filter(code=program.code, school_name=user.school_name).exclude(id=program.id)
        if existing.exists():
            return JsonResponse({'success': False, 'message': f'Program with code {program.code} already exists in your school.'})
        
        program.save()
        
        return JsonResponse({
            'success': True,
            'message': 'Program updated successfully.',
            'program': {
                'id': program.id,
                'code': program.code,
                'name': program.name,
                'department': program.department,
            }
        })
    except Exception as e:
        return JsonResponse({'success': False, 'message': str(e)})

@login_required
@admin_required
@require_http_methods(["POST"])
def admin_add_department_view(request):
    """Add a new department"""
    user = request.user
    
    try:
        # Debug logging
        logger.debug(f"Request method: {request.method}")
        logger.debug(f"Content-Type: {request.META.get('CONTENT_TYPE', '')}")
        logger.debug(f"POST data: {dict(request.POST)}")
        logger.debug(f"FILES data: {list(request.FILES.keys())}")
        
        # Handle both JSON and form data (FormData for file uploads)
        # Check if it's a JSON request by looking at content type
        content_type = request.META.get('CONTENT_TYPE', '')
        
        if 'application/json' in content_type:
            # JSON request
            data = json.loads(request.body)
            name = data.get('name', '').strip()
            code = data.get('code', '').strip().upper()
            icon = None
        else:
            # FormData request (for file uploads) - use POST data
            name = request.POST.get('name', '').strip()
            code = request.POST.get('code', '').strip().upper()
            # Handle icon upload
            icon = None
            if 'icon' in request.FILES:
                icon = request.FILES['icon']
                logger.debug(f"Icon file received: {icon.name}, size: {icon.size}")
        
        if not name:
            return JsonResponse({'success': False, 'message': 'Department name is required.'})
        
        if not code:
            return JsonResponse({'success': False, 'message': 'Department code is required.'})
        
        # Check if department with same code exists in this school (only if code is provided)
        if code:
            existing_code = Department.objects.filter(code=code, school_name=user.school_name)
            if existing_code.exists():
                return JsonResponse({'success': False, 'message': f'Department with code "{code}" already exists in your school.'})
        
        # Check if department with same name exists in this school
        existing_name = Department.objects.filter(name=name, school_name=user.school_name)
        if existing_name.exists():
            return JsonResponse({'success': False, 'message': f'Department "{name}" already exists in your school.'})
        
        try:
            department = Department.objects.create(
                name=name,
                code=code,
                icon=icon,
                school_name=user.school_name,
                education_level=user.education_level or 'university_college',
                is_active=True
            )
            logger.info(f"Department created successfully: {department.id} - {department.code} - {department.name}")
            
            return JsonResponse({
                'success': True,
                'message': 'Department created successfully.',
                'department': {
                    'id': department.id,
                    'name': department.name,
                    'code': department.code,
                }
            })
        except Exception as create_error:
            logger.error(f"Error creating department: {str(create_error)}")
            import traceback
            error_trace = traceback.format_exc()
            logger.error(f"Traceback: {error_trace}")
            return JsonResponse({'success': False, 'message': f'Error creating department: {str(create_error)}'})
    except Exception as e:
        import traceback
        from django.conf import settings
        error_trace = traceback.format_exc()
        logger.error(f"Error adding department: {str(e)}\n{error_trace}")
        # Return more detailed error in development
        error_message = str(e)
        if settings.DEBUG:
            error_message = f"{error_message}"
        return JsonResponse({'success': False, 'message': f'Error: {error_message}'})

@login_required
@admin_required
@require_http_methods(["POST"])
def admin_update_department_view(request, department_id):
    """Update an existing department"""
    user = request.user
    
    try:
        department = get_object_or_404(Department, id=department_id)
        
        # Verify the department belongs to the admin's school
        if department.school_name and department.school_name != user.school_name:
            return JsonResponse({'success': False, 'message': 'You do not have permission to edit this department.'})
        
        # Handle both JSON and form data
        content_type = request.META.get('CONTENT_TYPE', '')
        
        if 'application/json' in content_type:
            data = json.loads(request.body)
            department.name = data.get('name', department.name).strip()
            department.code = data.get('code', department.code).strip().upper()
            icon = None
        else:
            department.name = request.POST.get('name', department.name).strip()
            department.code = request.POST.get('code', department.code).strip().upper()
            # Handle icon upload
            if 'icon' in request.FILES:
                department.icon = request.FILES['icon']
        
        # Check for duplicate code in same school
        if department.code:
            existing_code = Department.objects.filter(code=department.code, school_name=user.school_name).exclude(id=department.id)
            if existing_code.exists():
                return JsonResponse({'success': False, 'message': f'Department with code "{department.code}" already exists in your school.'})
        
        # Check for duplicate name in same school
        existing_name = Department.objects.filter(name=department.name, school_name=user.school_name).exclude(id=department.id)
        if existing_name.exists():
            return JsonResponse({'success': False, 'message': f'Department "{department.name}" already exists in your school.'})
        
        department.save()
        
        return JsonResponse({
            'success': True,
            'message': 'Department updated successfully.',
            'department': {
                'id': department.id,
                'name': department.name,
                'code': department.code,
            }
        })
    except Exception as e:
        return JsonResponse({'success': False, 'message': str(e)})

@login_required
@admin_required
@require_http_methods(["POST"])
def admin_delete_department_view(request, department_id):
    """Delete a department"""
    user = request.user
    
    try:
        department = get_object_or_404(Department, id=department_id)
        
        # Verify the department belongs to the admin's school
        if department.school_name and department.school_name != user.school_name:
            return JsonResponse({'success': False, 'message': 'You do not have permission to delete this department.'})
        
        # Check if department has active (non-deleted) programs
        programs_count = Program.objects.filter(
            department=department.name, 
            school_name=user.school_name,
            is_active=True,
            deleted_at__isnull=True
        ).count()
        if programs_count > 0:
            return JsonResponse({
                'success': False, 
                'message': f'Cannot delete department. It has {programs_count} program(s). Please remove all programs first.'
            })
        
        # Soft delete: set deleted_at timestamp
        from django.utils import timezone
        department.deleted_at = timezone.now()
        department.is_active = False
        department.save()
        
        return JsonResponse({
            'success': True,
            'message': 'Department moved to trash. It will be permanently deleted after 30 days.'
        })
    except Exception as e:
        return JsonResponse({'success': False, 'message': str(e)})

@login_required
@admin_required
@require_http_methods(["POST"])
def admin_delete_program_view(request, program_id):
    """Delete a program"""
    user = request.user
    
    try:
        program = get_object_or_404(Program, id=program_id)
        
        # Verify the program belongs to the admin's school
        if program.school_name and program.school_name != user.school_name:
            return JsonResponse({'success': False, 'message': 'You do not have permission to delete this program.'})
        
        # Prevent deletion if there are users still assigned to this program
        active_users = CustomUser.objects.filter(program=program, deleted_at__isnull=True)
        if user.school_name:
            active_users = active_users.filter(school_name=user.school_name)
        user_count = active_users.count()
        if user_count > 0:
            return JsonResponse({
                'success': False,
                'message': f'Cannot delete program. It still has {user_count} user{"s" if user_count != 1 else ""}. Please remove or reassign them first.'
            })
        
        # Soft delete: set deleted_at timestamp
        from django.utils import timezone
        program.deleted_at = timezone.now()
        program.is_active = False
        program.save()
        
        return JsonResponse({'success': True, 'message': 'Program moved to trash. It will be permanently deleted after 30 days.'})
    except Exception as e:
        return JsonResponse({'success': False, 'message': str(e)})

@login_required
@admin_required
def admin_user_management_view(request):
    """User management view with department grouping and search"""
    user = request.user
    
    # Get search query and filters
    search_query = request.GET.get('search', '').strip()
    department_filter = request.GET.get('department_filter', '').strip()
    section_filter = request.GET.get('section_filter', '').strip()
    year_level_filter = request.GET.get('year_level_filter', '').strip()
    program_id = request.GET.get('program', '').strip()
    user_type = request.GET.get('type', '').strip()  # 'instructor' or 'student'
    
    # Get program if provided
    selected_program = None
    if program_id:
        try:
            if user.school_name:
                selected_program = Program.objects.get(id=program_id, school_name=user.school_name)
            else:
                selected_program = Program.objects.get(id=program_id)
        except Program.DoesNotExist:
            pass
    
    # Base querysets filtered by school (exclude deleted)
    if user.school_name:
        teachers_qs = CustomUser.objects.filter(is_teacher=True, school_name=user.school_name, deleted_at__isnull=True)
        students_qs = CustomUser.objects.filter(is_student=True, school_name=user.school_name, deleted_at__isnull=True)
        pending_users = CustomUser.objects.filter(
            is_teacher=True,
            is_approved=False,
            school_name=user.school_name,
            deleted_at__isnull=True
        ).order_by('-date_joined')
    else:
        teachers_qs = CustomUser.objects.filter(is_teacher=True, deleted_at__isnull=True)
        students_qs = CustomUser.objects.filter(is_student=True, deleted_at__isnull=True)
        pending_users = CustomUser.objects.filter(is_teacher=True, is_approved=False, deleted_at__isnull=True).order_by('-date_joined')
    
    # Apply department filter if provided
    if department_filter:
        teachers_qs = teachers_qs.filter(department=department_filter)
        students_qs = students_qs.filter(department=department_filter)
    
    # Apply section filter if provided (for students)
    if section_filter:
        students_qs = students_qs.filter(section=section_filter)
    
    # Apply year level filter if provided (for students)
    if year_level_filter:
        try:
            year_level = int(year_level_filter)
            students_qs = students_qs.filter(year_level=year_level)
        except ValueError:
            pass
    
    # Apply program filter if provided
    if selected_program:
        teachers_qs = teachers_qs.filter(program=selected_program)
        students_qs = students_qs.filter(program=selected_program)
    
    # Apply search filter if provided
    if search_query:
        teachers_qs = teachers_qs.filter(
            models.Q(full_name__icontains=search_query) |
            models.Q(email__icontains=search_query) |
            models.Q(school_id__icontains=search_query) |
            models.Q(department__icontains=search_query) |
            models.Q(username__icontains=search_query)
        )
        students_qs = students_qs.filter(
            models.Q(full_name__icontains=search_query) |
            models.Q(email__icontains=search_query) |
            models.Q(school_id__icontains=search_query) |
            models.Q(department__icontains=search_query) |
            models.Q(username__icontains=search_query)
        )
    
    # Group teachers by department, then by program
    # Order by creation date - newest appears right after the previous added
    teachers_by_dept = {}
    for teacher in teachers_qs.order_by('date_joined', 'id'):
        dept = teacher.department or 'Uncategorized'
        if dept not in teachers_by_dept:
            teachers_by_dept[dept] = {}
        program_key = teacher.program.code if teacher.program else 'No Program'
        if program_key not in teachers_by_dept[dept]:
            teachers_by_dept[dept][program_key] = []
        teachers_by_dept[dept][program_key].append(teacher)
    
    # Group students by department, then by program - order by creation date
    students_by_dept = {}
    for student in students_qs.order_by('date_joined', 'id'):
        dept = student.department or 'Uncategorized'
        if dept not in students_by_dept:
            students_by_dept[dept] = {}
        program_key = student.program.code if student.program else 'No Program'
        if program_key not in students_by_dept[dept]:
            students_by_dept[dept][program_key] = []
        students_by_dept[dept][program_key].append(student)
    
    # Get unique sections and year levels for filter dropdown (from students)
    if user.school_name:
        all_students_for_filters = CustomUser.objects.filter(is_student=True, school_name=user.school_name)
    else:
        all_students_for_filters = CustomUser.objects.filter(is_student=True)
    unique_sections = sorted(set([s.section for s in all_students_for_filters if s.section]))
    unique_year_levels = sorted(set([s.year_level for s in all_students_for_filters if s.year_level]))
    
    # Get all Department objects (like program management does) - exclude deleted, newest appears right after the previous added
    if user.school_name:
        all_departments_objects = Department.objects.filter(school_name=user.school_name, is_active=True, deleted_at__isnull=True).order_by('created_at', 'id')
    else:
        all_departments_objects = Department.objects.filter(is_active=True, deleted_at__isnull=True).order_by('created_at', 'id')
    
    # Apply education level filter BEFORE converting to list
    if user.education_level:
        all_departments_objects = all_departments_objects.filter(education_level=user.education_level)
    
    # Convert to list - newest is already at the end, which is correct
    all_departments_objects = list(all_departments_objects)
    
    # Get department names for filter dropdown (from Department model)
    all_departments = [dept.name for dept in all_departments_objects]
    
    # Reorganize users by Department objects (matching by department name)
    # This ensures changes in program management reflect in user management
    teachers_by_dept_obj = []
    for dept_obj in all_departments_objects:
        dept_teachers = {}
        for teacher in teachers_qs.filter(department=dept_obj.name):
            program_key = teacher.program.code if teacher.program else 'No Program'
            if program_key not in dept_teachers:
                dept_teachers[program_key] = []
            dept_teachers[program_key].append(teacher)
        
        if dept_teachers:  # Only add if there are teachers in this department
            teachers_by_dept_obj.append({
                'department': dept_obj,
                'programs': dept_teachers
            })
    
    students_by_dept_obj = []
    for dept_obj in all_departments_objects:
        dept_students = {}
        for student in students_qs.filter(department=dept_obj.name):
            program_key = student.program.code if student.program else 'No Program'
            if program_key not in dept_students:
                dept_students[program_key] = []
            dept_students[program_key].append(student)
        
        if dept_students:  # Only add if there are students in this department
            students_by_dept_obj.append({
                'department': dept_obj,
                'programs': dept_students
            })
    
    # Combine all users for "All Users" view
    all_users = list(teachers_qs) + list(students_qs)
    
    # Group all users by department, then by program (old structure for backward compatibility)
    all_users_by_dept = {}
    for user_obj in all_users:
        dept = user_obj.department or 'Uncategorized'
        if dept not in all_users_by_dept:
            all_users_by_dept[dept] = {}
        program_key = user_obj.program.code if user_obj.program else 'No Program'
        if program_key not in all_users_by_dept[dept]:
            all_users_by_dept[dept][program_key] = []
        all_users_by_dept[dept][program_key].append(user_obj)
    
    # Combine all users for "All Users" view, organized by Department objects
    # Show all departments that have programs, even if they don't have users yet
    all_users_by_dept_obj = []
    for dept_obj in all_departments_objects:
        dept_users = {}
        for user_obj in all_users:
            if user_obj.department == dept_obj.name:
                program_key = user_obj.program.code if user_obj.program else 'No Program'
                if program_key not in dept_users:
                    dept_users[program_key] = []
                dept_users[program_key].append(user_obj)
        
        # Get programs for this department to check if department should be shown
        # Exclude deleted programs (soft-deleted programs have deleted_at set)
        if user.school_name:
            dept_programs = Program.objects.filter(
                department=dept_obj.name,
                school_name=user.school_name,
                is_active=True,
                deleted_at__isnull=True
            ).order_by('created_at', 'id')
        else:
            dept_programs = Program.objects.filter(
                department=dept_obj.name,
                is_active=True,
                deleted_at__isnull=True
            ).order_by('created_at', 'id')
        
        if user.education_level:
            dept_programs = dept_programs.filter(education_level=user.education_level)
        
        # Convert to list - newest is already at the end, which is correct
        dept_programs = list(dept_programs)
        
        # Show all active departments, even if they have no programs yet
        # Build programs dict with program codes as keys
        programs_dict = {}
        for program in dept_programs:
            programs_dict[program.code] = dept_users.get(program.code, [])
        
        # Add department regardless of whether it has programs
        # This ensures departments are always visible in the database page
        all_users_by_dept_obj.append({
            'department': dept_obj,
            'programs': programs_dict
        })
    
    unread_notifications = AdminNotification.objects.filter(admin=user, is_read=False).count()
    
    # Get all programs for filter dropdown (grouped by department) - exclude deleted
    if user.school_name:
        all_programs = Program.objects.filter(school_name=user.school_name, deleted_at__isnull=True).order_by('code')
    else:
        all_programs = Program.objects.filter(deleted_at__isnull=True).order_by('code')
    
    if user.education_level:
        all_programs = all_programs.filter(education_level=user.education_level)
    
    # Group programs by department for filter dropdown
    all_programs_by_dept = {}
    for program in all_programs:
        dept = program.department or 'Uncategorized'
        if dept not in all_programs_by_dept:
            all_programs_by_dept[dept] = []
        all_programs_by_dept[dept].append(program)
    
    context = {
        'user': user,
        'teachers': teachers_qs,
        'students': students_qs,
        'all_users': all_users,
        'teachers_by_dept': teachers_by_dept,  # Keep for backward compatibility
        'students_by_dept': students_by_dept,  # Keep for backward compatibility
        'all_users_by_dept': all_users_by_dept,  # Keep for backward compatibility
        'teachers_by_dept_obj': teachers_by_dept_obj,  # New: Using Department objects
        'students_by_dept_obj': students_by_dept_obj,  # New: Using Department objects
        'all_users_by_dept_obj': all_users_by_dept_obj,  # New: Using Department objects
        'departments': all_departments_objects,  # Department objects for template
        'pending_users': pending_users,
        'unread_notifications': unread_notifications,
        'search_query': search_query,
        'all_departments': all_departments,  # Department names for filter dropdown
        'all_programs_by_dept': all_programs_by_dept,  # Programs grouped by department for filter
        'section_filter': section_filter,
        'year_level_filter': year_level_filter,
        'unique_sections': unique_sections,
        'unique_year_levels': unique_year_levels,
        'selected_program': selected_program,
        'user_type': user_type,
    }
    
    return render(request, 'dashboard/admin/admin_user_management.html', context)

@login_required
@admin_required
def admin_department_users_view(request, department_id):
    """View all programs in a specific department for user management"""
    user = request.user
    department = get_object_or_404(Department, id=department_id)
    
    # Verify department belongs to admin's school
    if user.school_name and department.school_name != user.school_name:
        return render(request, 'dashboard/shared/error.html', {
            'message': 'You are not authorized to view this department.'
        })
    
    # Get all programs in this department (exclude deleted) - first added stays first, newest appears second
    if user.school_name:
        programs = Program.objects.filter(
            department=department.name,
            school_name=user.school_name,
            is_active=True,
            deleted_at__isnull=True
        ).order_by('created_at', 'id')
    else:
        programs = Program.objects.filter(
            department=department.name,
            is_active=True,
            deleted_at__isnull=True
        ).order_by('created_at', 'id')
    
    # Convert to list - newest is already at the end, which is correct
    programs = list(programs)
    
    # Get program data with user counts
    program_data = []
    for program in programs:
        if user.school_name:
            teacher_count = CustomUser.objects.filter(
                program=program,
                school_name=user.school_name,
                is_teacher=True,
                is_approved=True
            ).count()
            student_count = CustomUser.objects.filter(
                program=program,
                school_name=user.school_name,
                is_student=True,
                is_approved=True
            ).count()
        else:
            teacher_count = CustomUser.objects.filter(
                program=program,
                is_teacher=True,
                is_approved=True
            ).count()
            student_count = CustomUser.objects.filter(
                program=program,
                is_student=True,
                is_approved=True
            ).count()
        
        total_users = teacher_count + student_count
        
        program_data.append({
            'program': program,
            'teacher_count': teacher_count,
            'student_count': student_count,
            'total_users': total_users
        })
    
    unread_notifications = AdminNotification.objects.filter(admin=user, is_read=False).count()
    
    context = {
        'user': user,
        'department': department,
        'program_data': program_data,
        'unread_notifications': unread_notifications,
    }
    
    return render(request, 'dashboard/admin/admin_department_users.html', context)

@login_required
@admin_required
def admin_notifications_view(request):
    """Notifications view"""
    user = request.user
    
    # Get all notifications for this admin
    notifications = AdminNotification.objects.filter(admin=user).order_by('-created_at')[:50]
    unread_count = AdminNotification.objects.filter(admin=user, is_read=False).count()
    
    context = {
        'user': user,
        'notifications': notifications,
        'unread_count': unread_count,
        'unread_notifications': unread_count,  # For sidebar
    }
    
    return render(request, 'dashboard/admin/admin_notifications.html', context)

@login_required
@admin_required
@require_http_methods(["POST"])
def admin_approve_user_view(request, user_id):
    """Approve a user - school_name is already set during signup"""
    from django.core.mail import send_mail
    from django.template.loader import render_to_string
    from django.conf import settings
    
    admin_user = request.user
    
    try:
        user = get_object_or_404(CustomUser, id=user_id)
        
        # Only approve if user is not already approved
        if user.is_approved:
            return JsonResponse({'success': False, 'message': 'User is already approved.'})
        
        # Verify user belongs to admin's school (if admin has a school)
        if admin_user.school_name and user.school_name != admin_user.school_name:
            return JsonResponse({
                'success': False,
                'message': f'This user belongs to {user.school_name}, not your school ({admin_user.school_name}).'
            })
        
        # Approve the user (school_name is already set from signup)
        user.is_approved = True
        user.save()
        
        # Mark related notifications as read
        AdminNotification.objects.filter(
            admin=admin_user,
            related_user=user,
            is_read=False
        ).update(is_read=True)
        
        # Create notification for the approved user (handle if table doesn't exist yet)
        try:
            from .models import UserNotification
            UserNotification.objects.create(
                user=user,
                notification_type='account_approved',
                title='Account Approved',
                message=f'Your account has been approved by the administrator. You can now access all features of the Attendance System.'
            )
        except Exception:
            # Table doesn't exist yet, skip notification creation
            pass
        
        # Send email notification to teacher if approved
        if user.is_teacher and user.email:
            try:
                subject = f'Account Approved - {user.school_name or "Attendance System"}'
                message = f"""
Dear {user.full_name or user.username},

Your teacher account has been approved by the administrator.

You can now log in to the Attendance System using your credentials:
- Email/ID: {user.email or user.school_id}
- School: {user.school_name or "N/A"}

Please log in at: {request.build_absolute_uri('/')}

If you have any questions, please contact your school administrator.

Best regards,
Attendance System Team
"""
                send_mail(
                    subject,
                    message,
                    settings.DEFAULT_FROM_EMAIL if hasattr(settings, 'DEFAULT_FROM_EMAIL') else 'noreply@attendancesystem.com',
                    [user.email],
                    fail_silently=True,  # Don't fail approval if email fails
                )
            except Exception as email_error:
                # Log email error but don't fail the approval
                import logging
                logger = logging.getLogger(__name__)
                logger.error(f"Failed to send approval email to {user.email}: {str(email_error)}")
        
        return JsonResponse({
            'success': True,
            'message': f'{user.full_name or user.username} has been approved.'
        })
    except Exception as e:
        return JsonResponse({'success': False, 'message': str(e)})

@login_required
@admin_required
@require_http_methods(["POST"])
def admin_reject_user_view(request, user_id):
    """Reject a user signup"""
    admin_user = request.user
    
    try:
        user = get_object_or_404(CustomUser, id=user_id)
        
        # Mark related notifications as read
        AdminNotification.objects.filter(
            admin=admin_user,
            related_user=user,
            is_read=False
        ).update(is_read=True)
        
        # Delete the user account
        user.delete()
        
        return JsonResponse({
            'success': True,
            'message': 'User signup has been rejected and account removed.'
        })
    except Exception as e:
        return JsonResponse({'success': False, 'message': str(e)})

@login_required
@admin_required
@require_http_methods(["POST"])
def admin_mark_notification_read_view(request, notification_id):
    """Mark a notification as read"""
    admin_user = request.user
    
    try:
        notification = get_object_or_404(AdminNotification, id=notification_id, admin=admin_user)
        notification.is_read = True
        notification.save()
        
        return JsonResponse({'success': True})
    except Exception as e:
        return JsonResponse({'success': False, 'message': str(e)})

@login_required
@admin_required
def admin_courses_view(request):
    """Course management view - shows departments (like academics page)"""
    user = request.user
    
    # Get search query
    search_query = request.GET.get('search', '').strip()
    department_filter = request.GET.get('department_filter', '').strip()
    
    # Get all departments - exclude deleted, newest appears right after the previous added
    if user.school_name:
        all_departments = Department.objects.filter(school_name=user.school_name, is_active=True, deleted_at__isnull=True).order_by('created_at', 'id')
    else:
        all_departments = Department.objects.filter(is_active=True, deleted_at__isnull=True).order_by('created_at', 'id')
    
    if user.education_level:
        all_departments = all_departments.filter(education_level=user.education_level)
    
    # Convert to list - newest is already at the end, which is correct
    all_departments = list(all_departments)
    
    # Get all programs to count courses per department (exclude deleted)
    if user.school_name:
        all_programs = Program.objects.filter(school_name=user.school_name, is_active=True, deleted_at__isnull=True)
    else:
        all_programs = Program.objects.filter(is_active=True, deleted_at__isnull=True)
    
    if user.education_level:
        all_programs = all_programs.filter(education_level=user.education_level)
    
    # Get courses to count per department
    if user.school_name:
        all_courses = Course.objects.filter(school_name=user.school_name)
    else:
        all_courses = Course.objects.all()
    
    if user.education_level:
        all_courses = all_courses.filter(program__education_level=user.education_level)
    
    # Group departments with program and course counts
    departments_with_data = []
    for dept in all_departments:
        # Filter programs by department name
        dept_programs = all_programs.filter(department=dept.name)
        program_count = dept_programs.count()
        
        # Count courses in this department
        course_count = 0
        for program in dept_programs:
            course_count += all_courses.filter(program=program).count()
        
        departments_with_data.append({
            'department': dept,
            'program_count': program_count,
            'course_count': course_count
        })
    
    # Apply search filter
    if search_query:
        departments_with_data = [
            dept_data for dept_data in departments_with_data
            if (search_query.lower() in dept_data['department'].name.lower() or
                (dept_data['department'].code and search_query.lower() in dept_data['department'].code.lower()))
        ]
    
    # Apply department filter
    if department_filter:
        departments_with_data = [
            dept_data for dept_data in departments_with_data
            if dept_data['department'].name == department_filter
        ]
    
    unread_notifications = AdminNotification.objects.filter(admin=user, is_read=False).count()
    
    context = {
        'user': user,
        'departments_with_data': departments_with_data,
        'all_departments': all_departments,
        'search_query': search_query,
        'department_filter': department_filter,
        'unread_notifications': unread_notifications,
    }
    
    return render(request, 'dashboard/admin/admin_courses.html', context)

@login_required
@admin_required
def admin_department_courses_view(request, department_id):
    """View all programs in a specific department for course management"""
    user = request.user
    department = get_object_or_404(Department, id=department_id)
    
    # Verify department belongs to admin's school
    if user.school_name and department.school_name != user.school_name:
        return render(request, 'dashboard/shared/error.html', {
            'message': 'You are not authorized to view this department.'
        })
    
    # Get search query
    search_query = request.GET.get('search', '').strip()
    
    # Get all programs in this department (exclude deleted) - newest appears right after the previous added
    if user.school_name:
        programs = Program.objects.filter(
            department=department.name,
            school_name=user.school_name,
            is_active=True,
            deleted_at__isnull=True
        ).order_by('created_at', 'id')
    else:
        programs = Program.objects.filter(
            department=department.name,
            is_active=True,
            deleted_at__isnull=True
        ).order_by('created_at', 'id')
    
    # Convert to list - newest is already at the end, which is correct
    programs = list(programs)
    
    # Apply search filter
    if search_query:
        # Filter the list using list comprehension
        programs = [p for p in programs if search_query.lower() in (p.code or '').lower() or search_query.lower() in (p.name or '').lower()]
    
    # Get program data with grouped course counts (collapse multi-section into one)
    program_data = []
    for program in programs:
        if user.school_name:
            courses_qs = Course.objects.filter(program=program, school_name=user.school_name)
        else:
            courses_qs = Course.objects.filter(program=program)
        # Group by code+name+semester+school_year
        keys = set()
        for c in courses_qs.values('code', 'name', 'semester', 'school_year'):
            norm_code = (c.get('code') or '').strip().upper()
            norm_name = (c.get('name') or '').strip().upper()
            norm_sem = (c.get('semester') or '').strip().lower()
            norm_sy = (c.get('school_year') or '').strip()
            keys.add((norm_code, norm_name, norm_sem, norm_sy))
        course_count = len(keys)
        
        program_data.append({
            'program': program,
            'course_count': course_count
        })
    
    unread_notifications = AdminNotification.objects.filter(admin=user, is_read=False).count()
    
    context = {
        'user': user,
        'department': department,
        'program_data': program_data,
        'search_query': search_query,
        'unread_notifications': unread_notifications,
    }
    
    return render(request, 'dashboard/admin/admin_department_courses.html', context)

@login_required
@admin_required
def admin_program_courses_view(request, program_id):
    """View all courses in a specific program"""
    user = request.user
    program = get_object_or_404(Program, id=program_id)
    
    # Verify program belongs to admin's school
    if user.school_name and program.school_name != user.school_name:
        return render(request, 'dashboard/shared/error.html', {
            'message': 'You are not authorized to view this program.'
        })
    
    # Get search query and filters
    search_query = request.GET.get('search', '').strip()
    year_level_filter = request.GET.get('year_level_filter', '').strip()
    semester_filter = request.GET.get('semester_filter', '').strip()
    
    # Get all courses in this program
    if user.school_name:
        courses = Course.objects.filter(program=program, school_name=user.school_name)
    else:
        courses = Course.objects.filter(program=program)
    
    # Apply filters
    if year_level_filter:
        try:
            courses = courses.filter(year_level=int(year_level_filter))
        except ValueError:
            pass
    
    if semester_filter:
        courses = courses.filter(semester=semester_filter)
    
    # Apply search filter
    if search_query:
        courses = courses.filter(
            models.Q(code__icontains=search_query) |
            models.Q(name__icontains=search_query) |
            models.Q(section__icontains=search_query)
        )
    
    courses = courses.order_by('year_level', 'semester', 'section', 'code')
    
    # Get unique year levels, semesters, and sections for filters
    unique_year_levels = sorted(set(courses.values_list('year_level', flat=True)))
    unique_semesters = Course.SEMESTER_CHOICES
    
    # Group courses by year level and semester, collapsing sections with same course
    courses_by_year_semester = {}
    for course in courses:
        year_key = f"Year {course.year_level}"
        semester_key = course.get_semester_display()
        section_key = course.section
        school_year = course.school_year or ''
        
        if year_key not in courses_by_year_semester:
            courses_by_year_semester[year_key] = {
                'semesters': {},
                'unique_semesters': [],
                'unique_school_years': [],
                'display_label': format_year_label(course.year_level)
            }
        year_entry = courses_by_year_semester[year_key]
        if semester_key not in year_entry['semesters']:
            year_entry['semesters'][semester_key] = {
                'courses': {}
            }
        semester_entry = year_entry['semesters'][semester_key]
        grouped_courses = semester_entry['courses']
        
        course_key = f"{(course.code or '').strip().lower()}|{(course.name or '').strip().lower()}|{school_year.lower() if isinstance(school_year, str) else school_year}"
        if course_key not in grouped_courses:
            grouped_courses[course_key] = {
                'course': course,
                'sections': set()
            }
        if section_key:
            grouped_courses[course_key]['sections'].add(section_key)
        
        # Track unique semesters, sections, and school years for this year
        if semester_key not in year_entry['unique_semesters']:
            year_entry['unique_semesters'].append(semester_key)
        if school_year and school_year not in year_entry['unique_school_years']:
            year_entry['unique_school_years'].append(school_year)
    
    # Sort unique semesters, sections, and school years for each year
    for year_key in courses_by_year_semester:
        year_entry = courses_by_year_semester[year_key]
        year_entry['unique_semesters'].sort()
        year_entry['unique_school_years'].sort(reverse=True)  # Most recent first
        
        for semester_key, semester_entry in year_entry['semesters'].items():
            course_groups = []
            for group in semester_entry['courses'].values():
                course_groups.append({
                    'course': group['course'],
                    'sections': sorted(group['sections'])
                })
            course_groups.sort(key=lambda item: (
                (item['course'].code or '').lower(),
                (item['course'].name or '').lower()
            ))
            semester_entry['course_groups'] = course_groups
            semester_entry.pop('courses', None)
    
    # Compute total grouped courses across all years/semesters (count each course once regardless of sections)
    total_grouped_courses = 0
    for year_entry in courses_by_year_semester.values():
        for sem_entry in year_entry.get('semesters', {}).values():
            total_grouped_courses += len(sem_entry.get('course_groups', []))
    
    # Get department object for navigation
    department = None
    if program.department:
        try:
            if user.school_name:
                department = Department.objects.get(name=program.department, school_name=user.school_name, is_active=True)
            else:
                department = Department.objects.filter(name=program.department, is_active=True).first()
        except Department.DoesNotExist:
            pass
        except Department.MultipleObjectsReturned:
            if user.school_name:
                department = Department.objects.filter(name=program.department, school_name=user.school_name, is_active=True).first()
            else:
                department = Department.objects.filter(name=program.department, is_active=True).first()
    
    # Get all programs for instructor selection (needed for add course modal)
    if user.school_name:
        all_programs = Program.objects.filter(school_name=user.school_name).order_by('code')
    else:
        all_programs = Program.objects.all().order_by('code')
    
    if user.education_level:
        all_programs = all_programs.filter(education_level=user.education_level)
    
    # Get all departments for instructor selection
    all_departments = sorted(set(all_programs.values_list('department', flat=True)))
    all_departments = [dept for dept in all_departments if dept]
    
    # Get teachers grouped by department and program
    if user.school_name:
        teachers = CustomUser.objects.filter(is_teacher=True, is_approved=True, school_name=user.school_name).order_by('department', 'program__code', 'full_name', 'username')
    else:
        teachers = CustomUser.objects.filter(is_teacher=True, is_approved=True).order_by('department', 'program__code', 'full_name', 'username')
    
    teachers_by_dept_program = {}
    for teacher in teachers:
        dept = teacher.department or 'Uncategorized'
        if dept not in teachers_by_dept_program:
            teachers_by_dept_program[dept] = {}
        program_key = teacher.program.code if teacher.program else 'No Program'
        if program_key not in teachers_by_dept_program[dept]:
            teachers_by_dept_program[dept][program_key] = []
        teachers_by_dept_program[dept][program_key].append({
            'id': teacher.id,
            'name': teacher.full_name or teacher.username,
            'email': teacher.email
        })
    
    # Generate school year options (current year ± 2 years)
    from datetime import datetime
    current_year = datetime.now().year
    school_year_options = []
    for year in range(current_year - 2, current_year + 3):
        school_year_options.append(f"{year}-{year+1}")
    
    unread_notifications = AdminNotification.objects.filter(admin=user, is_read=False).count()
    
    context = {
        'user': user,
        'program': program,
        'department': department,
        'courses': courses,
        'courses_by_year_semester': courses_by_year_semester,
        'total_grouped_courses': total_grouped_courses,
        'all_programs': all_programs,
        'all_departments': all_departments,
        'teachers_by_dept_program': teachers_by_dept_program,
        'unique_year_levels': unique_year_levels,
        'unique_semesters': unique_semesters,
        'school_year_options': school_year_options,
        'search_query': search_query,
        'year_level_filter': year_level_filter,
        'semester_filter': semester_filter,
        'unread_notifications': unread_notifications,
    }
    
    return render(request, 'dashboard/admin/admin_program_courses.html', context)

@login_required
@admin_required
@require_http_methods(["GET", "POST"])
def admin_add_course_view(request):
    """Add a new course - READ ONLY FOR ADMIN (instructors manage courses now)"""
    return JsonResponse({'success': False, 'message': 'Course management is now handled by instructors. Please contact an instructor to add courses.'}, status=403)

@login_required
@admin_required
@require_http_methods(["GET", "POST"])
def admin_update_course_view_old(request):
    """Update an existing course - OLD VERSION (kept for reference)"""
    user = request.user
    
    if request.method == 'POST':
        try:
            # Get form data
            code = request.POST.get('code', '').strip()
            name = request.POST.get('name', '').strip()
            program_id = request.POST.get('program', '').strip()
            year_level = request.POST.get('year_level', '').strip()
            section = request.POST.get('section', '').strip()
            semester = request.POST.get('semester', '1st').strip()
            school_year = request.POST.get('school_year', '').strip()
            instructor_id = request.POST.get('instructor', '').strip() or None
            room = request.POST.get('room', '').strip() or None
            days = request.POST.get('days', '').strip()
            start_time = request.POST.get('start_time', '').strip()
            end_time = request.POST.get('end_time', '').strip()
            
            # Validate required fields first
            if not all([code, name, program_id, year_level, section, days, start_time, end_time]):
                return JsonResponse({'success': False, 'message': 'Please fill in all required fields.'})
            
            # Get program (needed for color assignment and validation)
            program = get_object_or_404(Program, id=int(program_id))
            
            # Verify program belongs to admin's school
            if user.school_name and program.school_name != user.school_name:
                return JsonResponse({'success': False, 'message': 'Invalid program selected.'})
            
            # Auto-assign color from predefined palette - ensure no repeats within same program/semester
            color_palette = [
                '#3C4770', '#6366F1', '#8B5CF6', '#EC4899', '#F43F5E',
                '#EF4444', '#F59E0B', '#10B981', '#06B6D4', '#3B82F6',
                '#A855F7', '#F97316', '#14B8A6', '#0EA5E9', '#84CC16'
            ]
            
            # Get existing colors for this program and semester (for new courses, check all existing)
            existing_courses = Course.objects.filter(
                program=program,
                semester=semester,
                school_year=school_year or None
            )
            
            used_colors = set(existing_courses.values_list('color', flat=True))
            available_colors = [c for c in color_palette if c not in used_colors]
            
            # If all colors are used, start over (but try to avoid immediate repeats)
            if not available_colors:
                available_colors = color_palette
            
            color = random.choice(available_colors)
            
            # Get instructor if provided
            instructor = None
            if instructor_id:
                instructor = get_object_or_404(CustomUser, id=int(instructor_id), is_teacher=True)
                if user.school_name and instructor.school_name != user.school_name:
                    return JsonResponse({'success': False, 'message': 'Invalid instructor selected.'})
            
            # Create course
            course = Course.objects.create(
                code=code,
                name=name,
                program=program,
                year_level=int(year_level),
                section=section,
                semester=semester,
                school_year=school_year or None,
                instructor=instructor,
                room=room,
                days=days,
                start_time=start_time,
                end_time=end_time,
                color=color,
                school_name=user.school_name,
                is_active=True
            )
            
            # Handle day-specific schedules if provided
            day_schedules = request.POST.get('day_schedules', '')
            if day_schedules:
                try:
                    schedules_data = json.loads(day_schedules)
                    for schedule_data in schedules_data:
                        CourseSchedule.objects.create(
                            course=course,
                            day=schedule_data.get('day'),
                            start_time=schedule_data.get('start_time'),
                            end_time=schedule_data.get('end_time'),
                            room=schedule_data.get('room') or None
                        )
                except (json.JSONDecodeError, KeyError, ValueError) as e:
                    logger.warning(f"Error parsing day schedules: {str(e)}")
                    # Continue without day-specific schedules
            
            return JsonResponse({'success': True, 'message': 'Course added successfully!', 'course_id': course.id})
        except Exception as e:
            logger.error(f"Error adding course: {str(e)}")
            return JsonResponse({'success': False, 'message': f'Error: {str(e)}'})
    
    # GET request - return form data
    # Get programs for dropdown
    if user.school_name:
        programs = Program.objects.filter(school_name=user.school_name).order_by('code')
    else:
        programs = Program.objects.all().order_by('code')
    
    if user.education_level:
        programs = programs.filter(education_level=user.education_level)
    
    # Get teachers for instructor dropdown
    if user.school_name:
        teachers = CustomUser.objects.filter(is_teacher=True, is_approved=True, school_name=user.school_name).order_by('full_name', 'username')
    else:
        teachers = CustomUser.objects.filter(is_teacher=True, is_approved=True).order_by('full_name', 'username')
    
    return JsonResponse({
        'success': True,
        'programs': [{'id': p.id, 'code': p.code, 'name': p.name} for p in programs],
        'teachers': [{'id': t.id, 'name': t.full_name or t.username, 'email': t.email} for t in teachers],
        'semesters': Course.SEMESTER_CHOICES,
    })

@login_required
@admin_required
@require_http_methods(["GET", "POST"])
def admin_update_course_view(request, course_id):
    """Update an existing course - READ ONLY FOR ADMIN (instructors manage courses now)"""
    return JsonResponse({'success': False, 'message': 'Course management is now handled by instructors. Please contact an instructor to edit courses.'}, status=403)

@login_required
@admin_required
@require_http_methods(["POST"])
def admin_delete_course_view_old(request, course_id):
    """Delete a course - OLD VERSION (kept for reference)"""
    user = request.user
    course = get_object_or_404(Course, id=course_id)
    
    # Verify course belongs to admin's school
    if user.school_name and course.school_name != user.school_name:
        return JsonResponse({'success': False, 'message': 'You do not have permission to delete this course.'})
    
    if request.method == 'POST':
        try:
            # Get form data
            code = request.POST.get('code', '').strip()
            name = request.POST.get('name', '').strip()
            program_id = request.POST.get('program', '').strip()
            year_level = request.POST.get('year_level', '').strip()
            section = request.POST.get('section', '').strip()
            semester = request.POST.get('semester', '1st').strip()
            school_year = request.POST.get('school_year', '').strip()
            instructor_id = request.POST.get('instructor', '').strip() or None
            room = request.POST.get('room', '').strip() or None
            days = request.POST.get('days', '').strip()
            start_time = request.POST.get('start_time', '').strip()
            end_time = request.POST.get('end_time', '').strip()
            # Get color from form (user can change it, or keep existing)
            color = request.POST.get('color', course.color if course.color else '#3C4770').strip()
            is_active = request.POST.get('is_active', 'false') == 'true'
            
            # Validate required fields
            if not all([code, name, program_id, year_level, section, days, start_time, end_time]):
                return JsonResponse({'success': False, 'message': 'Please fill in all required fields.'})
            
            # Get program
            program = get_object_or_404(Program, id=int(program_id))
            
            # When updating, if color is being changed, check for conflicts
            # But allow user's choice (they can override if needed)
            if color != course.color:
                # Check if new color conflicts with other courses in same program/semester
                conflicting = Course.objects.filter(
                    program=program,
                    semester=semester,
                    school_year=school_year or None,
                    color=color
                ).exclude(id=course.id).exists()
                # Note: We allow the conflict - user's choice takes precedence
            
            # Verify program belongs to admin's school
            if user.school_name and program.school_name != user.school_name:
                return JsonResponse({'success': False, 'message': 'Invalid program selected.'})
            
            # Get instructor if provided
            instructor = None
            if instructor_id:
                instructor = get_object_or_404(CustomUser, id=int(instructor_id), is_teacher=True)
                if user.school_name and instructor.school_name != user.school_name:
                    return JsonResponse({'success': False, 'message': 'Invalid instructor selected.'})
            
            # Update course
            course.code = code
            course.name = name
            course.program = program
            course.year_level = int(year_level)
            course.section = section
            course.semester = semester
            course.school_year = school_year or None
            course.instructor = instructor
            course.room = room
            course.days = days
            course.start_time = start_time
            course.end_time = end_time
            course.color = color
            course.is_active = is_active
            course.save()
            
            # Handle day-specific schedules if provided
            day_schedules = request.POST.get('day_schedules', '')
            if day_schedules:
                try:
                    schedules_data = json.loads(day_schedules)
                    # Delete existing schedules for this course
                    CourseSchedule.objects.filter(course=course).delete()
                    # Create new schedules
                    for schedule_data in schedules_data:
                        CourseSchedule.objects.create(
                            course=course,
                            day=schedule_data.get('day'),
                            start_time=schedule_data.get('start_time'),
                            end_time=schedule_data.get('end_time'),
                            room=schedule_data.get('room') or None
                        )
                except (json.JSONDecodeError, KeyError, ValueError) as e:
                    logger.warning(f"Error parsing day schedules: {str(e)}")
                    # Continue without updating day-specific schedules
            else:
                # If no day_schedules provided, delete existing day-specific schedules
                CourseSchedule.objects.filter(course=course).delete()
            
            return JsonResponse({'success': True, 'message': 'Course updated successfully!'})
        except Exception as e:
            logger.error(f"Error updating course: {str(e)}")
            return JsonResponse({'success': False, 'message': f'Error: {str(e)}'})
    
    # GET request - return course data
    # Get programs for dropdown
    if user.school_name:
        programs = Program.objects.filter(school_name=user.school_name).order_by('code')
    else:
        programs = Program.objects.all().order_by('code')
    
    if user.education_level:
        programs = programs.filter(education_level=user.education_level)
    
    # Get teachers for instructor dropdown
    if user.school_name:
        teachers = CustomUser.objects.filter(is_teacher=True, is_approved=True, school_name=user.school_name).order_by('full_name', 'username')
    else:
        teachers = CustomUser.objects.filter(is_teacher=True, is_approved=True).order_by('full_name', 'username')
    
    # Get day-specific schedules if they exist
    schedules = course.course_schedules.all().order_by('day_order')
    day_schedules = []
    if schedules.exists():
        for schedule in schedules:
            day_schedules.append({
                'day': schedule.day,
                'start_time': str(schedule.start_time),
                'end_time': str(schedule.end_time),
                'room': schedule.room or ''
            })
    
    course_data = {
        'id': course.id,
        'code': course.code,
        'name': course.name,
        'program_id': course.program.id,
        'year_level': course.year_level,
        'section': course.section,
        'semester': course.semester,
        'school_year': course.school_year or '',
        'instructor_id': course.instructor.id if course.instructor else None,
        'instructor_dept': course.instructor.department if course.instructor and course.instructor.department else None,
        'instructor_program': course.instructor.program.code if course.instructor and course.instructor.program else None,
        'room': course.room or '',
        'days': course.days,
        'start_time': str(course.start_time),
        'end_time': str(course.end_time),
        'color': course.color,
        'is_active': course.is_active,
        'day_schedules': day_schedules,
    }
    
    return JsonResponse({
        'success': True,
        'course': course_data,
        'programs': [{'id': p.id, 'code': p.code, 'name': p.name} for p in programs],
        'teachers': [{'id': t.id, 'name': t.full_name or t.username, 'email': t.email} for t in teachers],
        'semesters': Course.SEMESTER_CHOICES,
    })

@login_required
@admin_required
@require_http_methods(["POST"])
def admin_delete_course_view(request, course_id):
    """Delete a course - READ ONLY FOR ADMIN (instructors manage courses now)"""
    return JsonResponse({'success': False, 'message': 'Course management is now handled by instructors. Please contact an instructor to delete courses.'}, status=403)

@login_required
@admin_required
@require_http_methods(["GET"])
def admin_course_detail_view(request, course_id):
    """Get course details for modal display"""
    try:
        course = get_object_or_404(Course, id=course_id)
        
        # Verify the course belongs to the admin's school
        if request.user.school_name and course.school_name != request.user.school_name:
            return JsonResponse({'success': False, 'message': 'Unauthorized'}, status=403)
        
        # Get course details with safe attribute access
        instructor_name = 'Not assigned'
        if course.instructor:
            full_name = course.instructor.full_name or course.instructor.username or 'N/A'
            email = course.instructor.email or 'N/A'
            instructor_name = f"{full_name} ({email})"
        
        # Note: program.department is a CharField (string), not a ForeignKey
        department_name = 'N/A'
        if course.program and course.program.department:
            department_name = course.program.department  # It's already a string
        
        program_name = 'N/A'
        if course.program:
            program_name = f"{course.program.code} - {course.program.name}"
        
        # Get day-specific schedules
        schedules = course.course_schedules.all().order_by('day_order')
        schedule_list = []
        if schedules.exists():
            for schedule in schedules:
                # Get attendance times - use schedule-specific if available, otherwise fall back to course-level
                attendance_start = schedule.attendance_start.strftime('%I:%M %p') if schedule.attendance_start else (course.attendance_start.strftime('%I:%M %p') if course.attendance_start else 'Not set')
                attendance_end = schedule.attendance_end.strftime('%I:%M %p') if schedule.attendance_end else (course.attendance_end.strftime('%I:%M %p') if course.attendance_end else 'Not set')
                
                schedule_list.append({
                    'day': schedule.get_day_display(),
                    'start_time': schedule.start_time.strftime('%I:%M %p'),
                    'end_time': schedule.end_time.strftime('%I:%M %p'),
                    'room': schedule.room or course.room or 'Not specified',
                    'attendance_start': attendance_start,
                    'attendance_end': attendance_end
                })
        
        course_data = {
            'id': course.id,
            'code': course.code or 'N/A',
            'name': course.name or 'N/A',
            'program': program_name,
            'department': department_name,
            'year_level': f"{course.year_level}{'st' if course.year_level == 1 else 'nd' if course.year_level == 2 else 'rd' if course.year_level == 3 else 'th'} Year",
            'section': course.section or 'N/A',
            'semester': course.get_semester_display() if course.semester else 'N/A',
            'school_year': course.school_year or 'N/A',
            'instructor': instructor_name,
            'room': course.room or 'Not specified',
            'days': course.days.replace(',', ', ') if course.days else 'N/A',
            'start_time': course.start_time.strftime('%I:%M %p') if course.start_time else 'N/A',
            'end_time': course.end_time.strftime('%I:%M %p') if course.end_time else 'N/A',
            'color': course.color or '#3C4770',
            'is_active': 'Active' if course.is_active else 'Inactive',
            'attendance_start': course.attendance_start.strftime('%I:%M %p') if course.attendance_start else 'Not set',
            'attendance_end': course.attendance_end.strftime('%I:%M %p') if course.attendance_end else 'Not set',
            'created_at': course.created_at.strftime('%B %d, %Y at %I:%M %p') if course.created_at else 'N/A',
            'updated_at': course.updated_at.strftime('%B %d, %Y at %I:%M %p') if course.updated_at else 'N/A',
            'day_schedules': schedule_list,  # List of day-specific schedules
            'has_day_schedules': len(schedule_list) > 0,  # Flag to indicate if day-specific schedules exist
        }
        
        # Sibling sections (same course identity across sections) for dropdown
        siblings_qs = Course.objects.filter(
            school_name=course.school_name,
            code=course.code,
            name=course.name,
            semester=course.semester,
            school_year=course.school_year,
            is_active=True
        ).order_by('section', 'id')
        sibling_sections = [
            {'id': c.id, 'section': (c.section or '').upper() or 'N/A'}
            for c in siblings_qs
        ]
        course_data['sibling_sections'] = sibling_sections
        
        return JsonResponse({'success': True, 'course': course_data})
    except Exception as e:
        logger.error(f"Error fetching course details: {str(e)}")
        return JsonResponse({'success': False, 'message': f'Error loading course details: {str(e)}'}, status=500)

@login_required
@admin_required
@require_http_methods(["POST"])
def admin_update_profile_view(request):
    """Update admin profile"""
    user = request.user
    
    try:
        # Update email if provided
        new_email = request.POST.get('email', '').strip().lower()
        if new_email:
            # Check if email is already taken by another user
            if CustomUser.objects.filter(email=new_email).exclude(id=user.id).exists():
                return JsonResponse({'success': False, 'message': 'This email is already in use by another account.'})
            user.email = new_email
        
        # Update school ID if provided
        new_school_id = request.POST.get('school_id', '').strip()
        if new_school_id:
            # Check if school ID is already taken by another user
            if CustomUser.objects.filter(school_id=new_school_id).exclude(id=user.id).exists():
                return JsonResponse({'success': False, 'message': 'This school ID is already in use by another account.'})
            user.school_id = new_school_id
        
        # Update custom password if provided
        new_password = request.POST.get('new_password', '').strip()
        confirm_password = request.POST.get('confirm_new_password', '').strip()
        current_password = request.POST.get('current_password', '').strip()
        
        if new_password:
            # Validate password length
            if len(new_password) < 6:
                return JsonResponse({'success': False, 'message': 'Password must be at least 6 characters long.'})
            
            # Check if passwords match
            if new_password != confirm_password:
                return JsonResponse({'success': False, 'message': 'New password and confirm password do not match.'})
            
            # If user already has a custom password, require current password
            if user.custom_password:
                if not current_password:
                    return JsonResponse({'success': False, 'message': 'Current password is required to change your password.'})
                # Verify current password (either temporary or custom)
                from django.contrib.auth import authenticate
                authenticated = authenticate(request, username=user.username, password=current_password)
                if not authenticated and not user.check_custom_password(current_password):
                    return JsonResponse({'success': False, 'message': 'Current password is incorrect.'})
            # If no custom password exists, they can add one without current password (first time)
            
            user.set_custom_password(new_password)
            # Also update the main password so admin can login with either password
            user.set_password(new_password)
            logger.info(f"Custom password {'updated' if user.custom_password else 'added'} for admin {user.username}")
        
        # Update school name
        school_name = request.POST.get('school_name', '').strip()
        if school_name:
            user.school_name = school_name
        elif school_name == '':
            # Allow clearing school name
            user.school_name = None
        
        # Update profile picture
        if 'profile_picture' in request.FILES:
            profile_picture = request.FILES['profile_picture']
            # Validate file size (max 5MB)
            if profile_picture.size > 5 * 1024 * 1024:
                return JsonResponse({'success': False, 'message': 'Profile picture size must be less than 5MB.'})
            # Validate file type
            if not profile_picture.content_type.startswith('image/'):
                return JsonResponse({'success': False, 'message': 'Please upload a valid image file.'})
            user.profile_picture = profile_picture
        
        user.save()
        return JsonResponse({'success': True, 'message': 'Profile updated successfully!'})
    except Exception as e:
        logger.error(f"Error updating profile: {str(e)}")
        return JsonResponse({'success': False, 'message': f'Error: {str(e)}'})

@login_required
@admin_required
def admin_reports_view(request):
    """Attendance reports view"""
    user = request.user
    unread_notifications = AdminNotification.objects.filter(admin=user, is_read=False).count()
    return render(request, 'dashboard/admin/admin_reports.html', {
        'user': user,
        'unread_notifications': unread_notifications,
    })

@login_required
@admin_required
@require_http_methods(["GET", "POST"])
def admin_add_teacher_view(request):
    """Add a new teacher"""
    admin_user = request.user
    
    if request.method == 'POST':
        form = AdminAddTeacherForm(request.POST, admin_user=admin_user)
        if form.is_valid():
            try:
                user, temp_password = form.save(admin_user)
                
                # Send email with credentials - try synchronously first to catch errors
                email_sent = False
                email_error = None
                user_email = user.email
                user_full_name = user.full_name
                user_school_id = user.school_id
                school_name = admin_user.school_name or 'the Attendance System'
                login_url = request.build_absolute_uri('/accounts/login-signup/')
                
                try:
                    email_subject = f"Welcome to {school_name} - Your Instructor Account"
                    email_message = f"""
Hello {user_full_name},

Welcome to {school_name}! Your instructor account has been successfully created by the school administrator.

Your login credentials:
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Email/Username: {user_email}
School ID: {user_school_id}
Temporary Password: {temp_password}
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

IMPORTANT SECURITY INFORMATION:
• This is your temporary password. Please change it immediately after logging in.
• You can log in using either your email ({user_email}) or your School ID ({user_school_id}).
• After logging in, go to your profile settings to set your own custom password.
• You can use either your temporary password or your custom password to log in.

Login at: {login_url}

If you have any questions or need assistance, please contact your school administrator.

Best regards,
{school_name} Administration Team
"""
                    # Try to send email synchronously to catch errors immediately
                    result = send_mail(
                        email_subject,
                        email_message,
                        settings.DEFAULT_FROM_EMAIL,
                        [user_email],
                        fail_silently=False
                    )
                    if result:
                        email_sent = True
                        logger.info(f"Welcome email sent successfully to {user_email}")
                    else:
                        email_error = "Email send returned False - check email configuration"
                        logger.warning(f"Email send returned False for {user_email} - check email configuration")
                except Exception as e:
                    email_error = str(e)
                    logger.error(f"Failed to send email to {user_email}: {str(e)}")
                    # Log more details about the error
                    import traceback
                    logger.error(f"Email error traceback: {traceback.format_exc()}")
                    
                    # Provide more helpful error messages
                    if 'BadCredentials' in email_error or 'Username and Password not accepted' in email_error or 'authentication failed' in email_error.lower():
                        email_error = 'Email server authentication failed. Please check email settings in settings.py (EMAIL_HOST_USER and EMAIL_HOST_PASSWORD).'
                    elif '535' in email_error or '5.7.0' in email_error:
                        email_error = 'Email authentication failed. Please ensure the Gmail App Password is correct and up to date.'
                    elif 'Connection refused' in email_error or 'timeout' in email_error.lower():
                        email_error = 'Unable to connect to email server. Please check your internet connection and email server settings.'
                
                # Build response message based on email status
                if email_sent:
                    success_message = f'Teacher "{user.full_name}" has been added successfully. A welcome email with login credentials has been sent to {user.email}.'
                else:
                    success_message = f'Teacher "{user.full_name}" has been added successfully. However, the welcome email could not be sent to {user.email}.'
                    if email_error:
                        success_message += f' Error: {email_error}'
                    success_message += f' Please manually provide the login credentials: Email: {user_email}, School ID: {user_school_id}, Temporary Password: {temp_password}'
                
                return JsonResponse({
                    'success': True,
                    'message': success_message,
                    'user_id': user.id,
                    'temp_password': temp_password,
                    'email_sent': email_sent,
                    'email_error': email_error if not email_sent else None
                })
            except Exception as e:
                logger.error(f"Error adding teacher: {str(e)}")
                return JsonResponse({
                    'success': False,
                    'message': f'Error adding teacher: {str(e)}'
                })
        else:
            errors = {field: errors[0] for field, errors in form.errors.items()}
            return JsonResponse({
                'success': False,
                'message': 'Please correct the errors below.',
                'errors': errors
            })
    
    # GET request - return form HTML
    form = AdminAddTeacherForm(admin_user=admin_user)
    programs = Program.objects.filter(
        school_name=admin_user.school_name,
        is_active=True,
        deleted_at__isnull=True
    ) if admin_user.school_name else Program.objects.filter(is_active=True, deleted_at__isnull=True)
    
    # Get program from query parameter if provided
    selected_program_id = request.GET.get('program', None)
    selected_program = None
    if selected_program_id:
        try:
            selected_program = Program.objects.get(id=selected_program_id, school_name=admin_user.school_name) if admin_user.school_name else Program.objects.get(id=selected_program_id)
        except Program.DoesNotExist:
            pass
    
    return render(request, 'dashboard/admin/admin_add_teacher.html', {
        'form': form,
        'programs': programs,
        'user': admin_user,
        'selected_program': selected_program,
        'unread_notifications': AdminNotification.objects.filter(admin=admin_user, is_read=False).count(),
    })

@login_required
@admin_required
@require_http_methods(["GET", "POST"])
def admin_add_student_view(request):
    """Add a new student"""
    admin_user = request.user
    
    if request.method == 'POST':
        form = AdminAddStudentForm(request.POST, admin_user=admin_user)
        if form.is_valid():
            try:
                user, temp_password = form.save(admin_user)
                
                # Send email with credentials - try synchronously first to catch errors
                email_sent = False
                email_error = None
                user_email = user.email
                user_full_name = user.full_name
                user_school_id = user.school_id
                school_name = admin_user.school_name or 'the Attendance System'
                login_url = request.build_absolute_uri('/accounts/login-signup/')
                
                try:
                    email_subject = f"Welcome to {school_name} - Your Student Account"
                    email_message = f"""
Hello {user_full_name},

Welcome to {school_name}! Your student account has been successfully created by the school administrator.

Your login credentials:
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Email/Username: {user_email}
School ID: {user_school_id}
Temporary Password: {temp_password}
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

IMPORTANT SECURITY INFORMATION:
• This is your temporary password. Please change it immediately after logging in.
• You can log in using either your email ({user_email}) or your School ID ({user_school_id}).
• After logging in, go to your profile settings to set your own custom password.
• You can use either your temporary password or your custom password to log in.

Login at: {login_url}

If you have any questions or need assistance, please contact your school administrator.

Best regards,
{school_name} Administration Team
"""
                    # Try to send email synchronously to catch errors immediately
                    result = send_mail(
                        email_subject,
                        email_message,
                        settings.DEFAULT_FROM_EMAIL,
                        [user_email],
                        fail_silently=False
                    )
                    if result:
                        email_sent = True
                        logger.info(f"Welcome email sent successfully to {user_email}")
                    else:
                        email_error = "Email send returned False - check email configuration"
                        logger.warning(f"Email send returned False for {user_email} - check email configuration")
                except Exception as e:
                    email_error = str(e)
                    logger.error(f"Failed to send email to {user_email}: {str(e)}")
                    # Log more details about the error
                    import traceback
                    logger.error(f"Email error traceback: {traceback.format_exc()}")
                    
                    # Provide more helpful error messages
                    if 'BadCredentials' in email_error or 'Username and Password not accepted' in email_error or 'authentication failed' in email_error.lower():
                        email_error = 'Email server authentication failed. Please check email settings in settings.py (EMAIL_HOST_USER and EMAIL_HOST_PASSWORD).'
                    elif '535' in email_error or '5.7.0' in email_error:
                        email_error = 'Email authentication failed. Please ensure the Gmail App Password is correct and up to date.'
                    elif 'Connection refused' in email_error or 'timeout' in email_error.lower():
                        email_error = 'Unable to connect to email server. Please check your internet connection and email server settings.'
                
                # Build response message based on email status
                if email_sent:
                    success_message = f'Student "{user.full_name}" has been added successfully. A welcome email with login credentials has been sent to {user.email}.'
                else:
                    success_message = f'Student "{user.full_name}" has been added successfully. However, the welcome email could not be sent to {user.email}.'
                    if email_error:
                        success_message += f' Error: {email_error}'
                    success_message += f' Please manually provide the login credentials: Email: {user_email}, School ID: {user_school_id}, Temporary Password: {temp_password}'
                
                return JsonResponse({
                    'success': True,
                    'message': success_message,
                    'user_id': user.id,
                    'temp_password': temp_password,
                    'email_sent': email_sent,
                    'email_error': email_error if not email_sent else None
                })
            except Exception as e:
                logger.error(f"Error adding student: {str(e)}")
                return JsonResponse({
                    'success': False,
                    'message': f'Error adding student: {str(e)}'
                })
        else:
            errors = {field: errors[0] for field, errors in form.errors.items()}
            return JsonResponse({
                'success': False,
                'message': 'Please correct the errors below.',
                'errors': errors
            })
    
    # GET request - return form HTML
    form = AdminAddStudentForm(admin_user=admin_user)
    programs = Program.objects.filter(
        school_name=admin_user.school_name,
        is_active=True,
        deleted_at__isnull=True
    ) if admin_user.school_name else Program.objects.filter(is_active=True, deleted_at__isnull=True)
    
    # Get program from query parameter if provided
    selected_program_id = request.GET.get('program', None)
    selected_program = None
    if selected_program_id:
        try:
            selected_program = Program.objects.get(id=selected_program_id, school_name=admin_user.school_name) if admin_user.school_name else Program.objects.get(id=selected_program_id)
        except Program.DoesNotExist:
            pass
    
    return render(request, 'dashboard/admin/admin_add_student.html', {
        'form': form,
        'programs': programs,
        'user': admin_user,
        'selected_program': selected_program,
        'unread_notifications': AdminNotification.objects.filter(admin=admin_user, is_read=False).count(),
    })

@login_required
@admin_required
def admin_user_detail_view(request, user_id):
    """Get user details for modal display"""
    user = get_object_or_404(CustomUser, id=user_id)
    
    # Verify the user belongs to the admin's school
    if request.user.school_name and user.school_name != request.user.school_name:
        return JsonResponse({'success': False, 'message': 'Unauthorized'}, status=403)
    
    # Get user details
    # Note: We can't retrieve the original password, but we can check if it was recently created
    user_data = {
        'id': user.id,
        'full_name': user.full_name or 'N/A',
        'email': user.email or 'N/A',
        'school_id': user.school_id or 'N/A',
        'username': user.username or 'N/A',
        'education_level': user.get_education_level_display() if user.education_level else 'N/A',
        'school_name': user.school_name or 'N/A',
        'department': user.department or 'N/A',
        'program': f"{user.program.code} - {user.program.name}" if user.program else 'N/A',
        'year_level': f"{user.year_level}{'st' if user.year_level == 1 else 'nd' if user.year_level == 2 else 'rd' if user.year_level == 3 else 'th'} Year" if user.year_level else 'N/A',
        'section': user.section or 'N/A',
        'is_approved': user.is_approved,
        'date_joined': (
            timezone.localtime(
                user.date_joined if not timezone.is_naive(user.date_joined) else timezone.make_aware(user.date_joined, timezone.get_default_timezone()),
                PH_TIMEZONE
            ).strftime('%B %d, %Y %I:%M %p') if user.date_joined else 'N/A'
        ),
        'user_type': 'Instructor' if user.is_teacher else 'Student',
        'profile_picture': user.profile_picture.url if user.profile_picture else '',
        'password': None,  # Password cannot be retrieved, but can be shown if passed via query param
    }
    
    # Check if password was passed in request (from recent creation) or retrieve from storage
    temp_password = request.GET.get('password', '')
    if not temp_password:
        # Try to retrieve the stored temporary password
        try:
            from .models import UserTemporaryPassword
            password_record = UserTemporaryPassword.objects.get(user=user)
            temp_password = password_record.password
        except UserTemporaryPassword.DoesNotExist:
            temp_password = None
    
    if temp_password:
        user_data['password'] = temp_password
    
    return JsonResponse({'success': True, 'user': user_data})

@login_required
@admin_required
@require_http_methods(["POST"])
def admin_delete_user_view(request, user_id):
    """Delete a user"""
    user = get_object_or_404(CustomUser, id=user_id)
    
    # Verify the user belongs to the admin's school
    if request.user.school_name and user.school_name != request.user.school_name:
        return JsonResponse({'success': False, 'message': 'Unauthorized'}, status=403)
    
    user_name = user.full_name or user.username
    user_type = 'Instructor' if user.is_teacher else 'Student'
    
    try:
        # Soft delete: set deleted_at timestamp
        from django.utils import timezone
        user.deleted_at = timezone.now()
        user.is_active = False
        user.save()
        return JsonResponse({'success': True, 'message': f'{user_type} "{user_name}" moved to trash. It will be permanently deleted after 30 days.'})
    except Exception as e:
        logger.error(f"Error deleting user: {str(e)}")
        return JsonResponse({'success': False, 'message': f'Error deleting user: {str(e)}'})


@login_required
@admin_required
@require_http_methods(["POST"])
def admin_bulk_delete_users_view(request, program_id):
    """Bulk delete users (database view only)."""
    user = request.user
    program = get_object_or_404(Program, id=program_id)
    
    if user.school_name and program.school_name != user.school_name:
        return JsonResponse({'success': False, 'message': 'You do not have permission to modify this program.'}, status=403)
    
    try:
        data = json.loads(request.body or '{}')
    except json.JSONDecodeError:
        data = {}
    
    scope = (data.get('scope') or '').lower()
    user_ids = data.get('user_ids') or []
    valid_scopes = {'all', 'instructors', 'students', 'specific'}
    if scope not in valid_scopes:
        return JsonResponse({'success': False, 'message': 'Invalid delete option selected.'}, status=400)
    
    queryset = CustomUser.objects.filter(program=program, deleted_at__isnull=True)
    if user.school_name:
        queryset = queryset.filter(school_name=user.school_name)
    
    if scope == 'specific':
        try:
            user_ids = [int(uid) for uid in user_ids if str(uid).isdigit()]
        except (ValueError, TypeError):
            user_ids = []
        if not user_ids:
            return JsonResponse({'success': False, 'message': 'Please select at least one user to delete.'}, status=400)
        queryset = queryset.filter(id__in=user_ids)
    elif scope == 'instructors':
        queryset = queryset.filter(is_teacher=True)
    elif scope == 'students':
        queryset = queryset.filter(is_student=True)
    else:
        queryset = queryset.filter(models.Q(is_teacher=True) | models.Q(is_student=True))
    
    count = queryset.count()
    if count == 0:
        return JsonResponse({'success': False, 'message': 'No users found for the selected option.'})
    
    from django.utils import timezone
    now = timezone.now()
    queryset.update(deleted_at=now, is_active=False)
    
    return JsonResponse({'success': True, 'message': f'{count} user{"s" if count != 1 else ""} removed successfully.'})

@login_required
@admin_required
def admin_download_user_document_view(request, user_id):
    """Download Word document with user information"""
    user = get_object_or_404(CustomUser, id=user_id)
    
    # Verify the user belongs to the admin's school
    if request.user.school_name and user.school_name != request.user.school_name:
        return HttpResponse('Unauthorized', status=403)
    
    # Get temporary password from request or retrieve from stored record
    temp_password = request.GET.get('password', '')
    if not temp_password:
        # Try to retrieve the original temporary password from storage
        try:
            from .models import UserTemporaryPassword
            password_record = UserTemporaryPassword.objects.get(user=user)
            temp_password = password_record.password
            logger.info(f"Retrieved stored temporary password for user {user.id}")
        except UserTemporaryPassword.DoesNotExist:
            # If no stored password exists (for users created before this feature),
            # we cannot retrieve the original password. Use a placeholder message.
            temp_password = 'Password was set during account creation. Please contact administrator for password reset if needed.'
            logger.warning(f"No stored temporary password found for user {user.id} - user may have been created before password storage feature")
    
    # Determine user type
    user_type = 'instructor' if user.is_teacher else 'student'
    
    # Generate document
    if not DOCX_AVAILABLE:
        logger.error("python-docx is not available. Please ensure it is installed: pip install python-docx")
        return HttpResponse('Word document generation is not available. Please install python-docx: pip install python-docx', status=500)
    
    try:
        # Build the login URL
        login_url = request.build_absolute_uri('/accounts/login-signup/')
        buffer = generate_user_document(user, temp_password, user_type, login_url)
        if not buffer:
            logger.error("generate_user_document returned None")
            return HttpResponse('Error generating document: Document generation returned None', status=500)
        
        # Create response
        filename = f"{user.full_name or user.username}_{user_type}_credentials.docx"
        # Sanitize filename - remove special characters
        import re
        filename = re.sub(r'[^\w\s-]', '', filename).strip()
        filename = re.sub(r'[-\s]+', '-', filename)
        
        response = HttpResponse(
            buffer.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        
        logger.info(f"Successfully generated document for user {user.id}: {filename}")
        return response
    except Exception as e:
        logger.error(f"Error generating document: {str(e)}", exc_info=True)
        return HttpResponse(f'Error generating document: {str(e)}. Please check server logs for details.', status=500)

# ============================================
# TRASH MANAGEMENT VIEWS
# ============================================

@login_required
@admin_required
def admin_trash_view(request):
    """View all deleted items (departments, programs, users) categorized"""
    user = request.user
    category = request.GET.get('category', 'all')  # 'all', 'departments', 'programs', 'users'
    
    from django.utils import timezone
    
    # Get deleted items
    deleted_departments = []
    deleted_programs = []
    deleted_users = []
    
    if user.school_name:
        deleted_departments = Department.objects.filter(
            deleted_at__isnull=False,
            school_name=user.school_name
        ).order_by('-deleted_at')
        
        deleted_programs = Program.objects.filter(
            deleted_at__isnull=False,
            school_name=user.school_name
        ).order_by('-deleted_at')
        
        deleted_users = CustomUser.objects.filter(
            deleted_at__isnull=False,
            school_name=user.school_name
        ).order_by('-deleted_at')
    else:
        deleted_departments = Department.objects.filter(deleted_at__isnull=False).order_by('-deleted_at')
        deleted_programs = Program.objects.filter(deleted_at__isnull=False).order_by('-deleted_at')
        deleted_users = CustomUser.objects.filter(deleted_at__isnull=False).order_by('-deleted_at')
    
    if user.education_level:
        deleted_departments = deleted_departments.filter(education_level=user.education_level)
        deleted_programs = deleted_programs.filter(education_level=user.education_level)
    
    # Calculate days until permanent deletion (30 days from deleted_at)
    now = timezone.now()
    for dept in deleted_departments:
        days_until_deletion = 30 - (now - dept.deleted_at).days
        dept.days_until_deletion = max(0, days_until_deletion)
    
    for prog in deleted_programs:
        days_until_deletion = 30 - (now - prog.deleted_at).days
        prog.days_until_deletion = max(0, days_until_deletion)
    
    for usr in deleted_users:
        days_until_deletion = 30 - (now - usr.deleted_at).days
        usr.days_until_deletion = max(0, days_until_deletion)
    
    # Filter by category
    if category == 'departments':
        deleted_programs = []
        deleted_users = []
    elif category == 'programs':
        deleted_departments = []
        deleted_users = []
    elif category == 'users':
        deleted_departments = []
        deleted_programs = []
    
    unread_notifications = AdminNotification.objects.filter(admin=user, is_read=False).count()
    
    context = {
        'user': user,
        'deleted_departments': deleted_departments,
        'deleted_programs': deleted_programs,
        'deleted_users': deleted_users,
        'category': category,
        'unread_notifications': unread_notifications,
    }
    
    return render(request, 'dashboard/admin/admin_trash.html', context)

@login_required
@admin_required
@require_http_methods(["POST"])
def admin_restore_department_view(request, department_id):
    """Restore a deleted department"""
    user = request.user
    
    try:
        department = get_object_or_404(Department, id=department_id, deleted_at__isnull=False)
        
        # Verify the department belongs to the admin's school
        if department.school_name and department.school_name != user.school_name:
            return JsonResponse({'success': False, 'message': 'You do not have permission to restore this department.'})
        
        # Restore: clear deleted_at and set is_active
        department.deleted_at = None
        department.is_active = True
        department.save()
        
        return JsonResponse({
            'success': True,
            'message': 'Department restored successfully.'
        })
    except Exception as e:
        return JsonResponse({'success': False, 'message': str(e)})

@login_required
@admin_required
@require_http_methods(["POST"])
def admin_restore_program_view(request, program_id):
    """Restore a deleted program"""
    user = request.user
    
    try:
        program = get_object_or_404(Program, id=program_id, deleted_at__isnull=False)
        
        # Verify the program belongs to the admin's school
        if program.school_name and program.school_name != user.school_name:
            return JsonResponse({'success': False, 'message': 'You do not have permission to restore this program.'})
        
        # Restore: clear deleted_at and set is_active
        program.deleted_at = None
        program.is_active = True
        program.save()
        
        return JsonResponse({'success': True, 'message': 'Program restored successfully.'})
    except Exception as e:
        return JsonResponse({'success': False, 'message': str(e)})

@login_required
@admin_required
@require_http_methods(["POST"])
def admin_restore_user_view(request, user_id):
    """Restore a deleted user"""
    user = request.user
    
    try:
        deleted_user = get_object_or_404(CustomUser, id=user_id, deleted_at__isnull=False)
        
        # Verify the user belongs to the admin's school
        if deleted_user.school_name and deleted_user.school_name != user.school_name:
            return JsonResponse({'success': False, 'message': 'You do not have permission to restore this user.'})
        
        # Restore: clear deleted_at and set is_active
        deleted_user.deleted_at = None
        deleted_user.is_active = True
        deleted_user.save()
        
        user_name = deleted_user.full_name or deleted_user.username
        user_type = 'Instructor' if deleted_user.is_teacher else 'Student'
        
        return JsonResponse({'success': True, 'message': f'{user_type} "{user_name}" restored successfully.'})
    except Exception as e:
        logger.error(f"Error restoring user: {str(e)}")
        return JsonResponse({'success': False, 'message': f'Error restoring user: {str(e)}'})

@login_required
@admin_required
@require_http_methods(["POST"])
def admin_permanent_delete_department_view(request, department_id):
    """Permanently delete a soft-deleted department"""
    user = request.user
    
    try:
        department = get_object_or_404(Department, id=department_id, deleted_at__isnull=False)
        
        # Verify the department belongs to the admin's school
        if department.school_name and department.school_name != user.school_name:
            return JsonResponse({'success': False, 'message': 'You do not have permission to delete this department.'})
        
        dept_name = department.code + ' - ' + department.name if department.code else department.name
        
        # Permanent delete
        department.delete()
        
        return JsonResponse({'success': True, 'message': f'Department "{dept_name}" permanently deleted.'})
    except Exception as e:
        logger.error(f"Error permanently deleting department: {str(e)}")
        return JsonResponse({'success': False, 'message': f'Error permanently deleting department: {str(e)}'})

@login_required
@admin_required
@require_http_methods(["POST"])
def admin_permanent_delete_program_view(request, program_id):
    """Permanently delete a soft-deleted program"""
    user = request.user
    
    try:
        program = get_object_or_404(Program, id=program_id, deleted_at__isnull=False)
        
        # Verify the program belongs to the admin's school
        if program.school_name and program.school_name != user.school_name:
            return JsonResponse({'success': False, 'message': 'You do not have permission to delete this program.'})
        
        prog_name = f'{program.code} - {program.name}'
        
        # Permanent delete
        program.delete()
        
        return JsonResponse({'success': True, 'message': f'Program "{prog_name}" permanently deleted.'})
    except Exception as e:
        logger.error(f"Error permanently deleting program: {str(e)}")
        return JsonResponse({'success': False, 'message': f'Error permanently deleting program: {str(e)}'})

@login_required
@admin_required
@require_http_methods(["POST"])
def admin_permanent_delete_user_view(request, user_id):
    """Permanently delete a soft-deleted user"""
    user = request.user
    
    try:
        deleted_user = get_object_or_404(CustomUser, id=user_id, deleted_at__isnull=False)
        
        # Verify the user belongs to the admin's school
        if deleted_user.school_name and deleted_user.school_name != user.school_name:
            return JsonResponse({'success': False, 'message': 'You do not have permission to delete this user.'})
        
        user_name = deleted_user.full_name or deleted_user.username
        user_type = 'Instructor' if deleted_user.is_teacher else 'Student'
        
        # Permanent delete
        deleted_user.delete()
        
        return JsonResponse({'success': True, 'message': f'{user_type} "{user_name}" permanently deleted.'})
    except Exception as e:
        logger.error(f"Error permanently deleting user: {str(e)}")
        return JsonResponse({'success': False, 'message': f'Error permanently deleting user: {str(e)}'})

@login_required
@admin_required
@require_http_methods(["POST"])
def admin_restore_all_trash_view(request):
    """Restore all items in trash"""
    user = request.user
    
    try:
        dept_count = 0
        prog_count = 0
        user_count = 0
        
        if user.school_name:
            deleted_departments = Department.objects.filter(
                deleted_at__isnull=False,
                school_name=user.school_name
            )
            deleted_programs = Program.objects.filter(
                deleted_at__isnull=False,
                school_name=user.school_name
            )
            deleted_users = CustomUser.objects.filter(
                deleted_at__isnull=False,
                school_name=user.school_name
            )
        else:
            deleted_departments = Department.objects.filter(deleted_at__isnull=False)
            deleted_programs = Program.objects.filter(deleted_at__isnull=False)
            deleted_users = CustomUser.objects.filter(deleted_at__isnull=False)
        
        if user.education_level:
            deleted_departments = deleted_departments.filter(education_level=user.education_level)
            deleted_programs = deleted_programs.filter(education_level=user.education_level)
        
        dept_count = deleted_departments.count()
        prog_count = deleted_programs.count()
        user_count = deleted_users.count()
        
        # Restore all items
        deleted_departments.update(deleted_at=None, is_active=True)
        deleted_programs.update(deleted_at=None, is_active=True)
        deleted_users.update(deleted_at=None, is_active=True)
        
        total_count = dept_count + prog_count + user_count
        message = f'Successfully restored {total_count} item(s) from trash'
        if dept_count > 0:
            message += f' ({dept_count} department{"s" if dept_count != 1 else ""}'
        if prog_count > 0:
            message += f'{", " if dept_count > 0 else " ("}{prog_count} program{"s" if prog_count != 1 else ""}'
        if user_count > 0:
            message += f'{", " if dept_count > 0 or prog_count > 0 else " ("}{user_count} user{"s" if user_count != 1 else ""}'
        if dept_count > 0 or prog_count > 0 or user_count > 0:
            message += ')'
        
        return JsonResponse({'success': True, 'message': message})
    except Exception as e:
        logger.error(f"Error restoring all trash items: {str(e)}")
        return JsonResponse({'success': False, 'message': f'Error restoring all items: {str(e)}'})

@login_required
@admin_required
@require_http_methods(["POST"])
def admin_delete_all_trash_view(request):
    """Permanently delete all items in trash"""
    user = request.user
    
    try:
        dept_count = 0
        prog_count = 0
        user_count = 0
        
        if user.school_name:
            deleted_departments = Department.objects.filter(
                deleted_at__isnull=False,
                school_name=user.school_name
            )
            deleted_programs = Program.objects.filter(
                deleted_at__isnull=False,
                school_name=user.school_name
            )
            deleted_users = CustomUser.objects.filter(
                deleted_at__isnull=False,
                school_name=user.school_name
            )
        else:
            deleted_departments = Department.objects.filter(deleted_at__isnull=False)
            deleted_programs = Program.objects.filter(deleted_at__isnull=False)
            deleted_users = CustomUser.objects.filter(deleted_at__isnull=False)
        
        if user.education_level:
            deleted_departments = deleted_departments.filter(education_level=user.education_level)
            deleted_programs = deleted_programs.filter(education_level=user.education_level)
        
        dept_count = deleted_departments.count()
        prog_count = deleted_programs.count()
        user_count = deleted_users.count()
        
        # Permanently delete all items
        deleted_departments.delete()
        deleted_programs.delete()
        deleted_users.delete()
        
        total_count = dept_count + prog_count + user_count
        message = f'Successfully deleted {total_count} item(s) from trash'
        if dept_count > 0:
            message += f' ({dept_count} department{"s" if dept_count != 1 else ""}'
        if prog_count > 0:
            message += f'{", " if dept_count > 0 else " ("}{prog_count} program{"s" if prog_count != 1 else ""}'
        if user_count > 0:
            message += f'{", " if dept_count > 0 or prog_count > 0 else " ("}{user_count} user{"s" if user_count != 1 else ""}'
        if dept_count > 0 or prog_count > 0 or user_count > 0:
            message += ')'
        
        return JsonResponse({'success': True, 'message': message})
    except Exception as e:
        logger.error(f"Error deleting all trash items: {str(e)}")
        return JsonResponse({'success': False, 'message': f'Error deleting all items: {str(e)}'})

